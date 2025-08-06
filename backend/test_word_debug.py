#!/usr/bin/env python3
"""Test Word export with minimal error catching"""

import os
import sys
sys.path.append('.')

try:
    print("Testing Word export step by step...")
    
    # Test 1: Import docx
    print("1. Testing python-docx import...")
    from docx import Document
    print("✅ python-docx imported successfully")
    
    # Test 2: Create minimal document
    print("2. Creating minimal document...")
    doc = Document()
    doc.add_heading('Test', 0)
    doc.add_paragraph('Test paragraph')
    
    # Test 3: Create table
    print("3. Creating table...")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = 'Header 1'
    table.rows[0].cells[1].text = 'Header 2'
    table.rows[1].cells[0].text = 'Data 1'
    table.rows[1].cells[1].text = 'Data 2'
    print("✅ Table created successfully")
    
    # Test 4: Save document
    print("4. Saving document...")
    test_file = os.path.join('exports', 'debug_test.docx')
    os.makedirs('exports', exist_ok=True)
    doc.save(test_file)
    print(f"✅ Document saved: {test_file}")
    
    # Test 5: Test with app context
    print("5. Testing with Flask app context...")
    from flask import Flask
    from flask_sqlalchemy import SQLAlchemy
    from models_new import init_models
    
    app = Flask(__name__)
    basedir = os.path.dirname(os.path.abspath(__file__))
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(basedir, "migration_tool.db")}'
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    
    db = SQLAlchemy(app)
    
    with app.app_context():
        models = init_models(db)
        print("✅ Flask app context and models initialized")
        
        # Test 6: Get summary data manually
        print("6. Testing summary data...")
        Server = models.get('Server')
        Database = models.get('Database') 
        FileShare = models.get('FileShare')
        
        servers_count = Server.query.count() if Server else 0
        databases_count = Database.query.count() if Database else 0
        file_shares_count = FileShare.query.count() if FileShare else 0
        
        print(f"Servers: {servers_count}, Databases: {databases_count}, File Shares: {file_shares_count}")
        
        # Test 7: Create metrics data like in Word export
        print("7. Testing metrics data structure...")
        metrics_data = [
            ('Total Infrastructure Components', str(servers_count + databases_count + file_shares_count)),
            ('Estimated Migration Duration', '12 weeks'),
            ('Migration Complexity', 'Medium'),
            ('Estimated Monthly Cloud Cost', '$5,000.00'),
            ('Estimated Annual Cloud Cost', '$60,000.00'),
            ('Primary Migration Strategy', 'Rehost (Lift & Shift)')
        ]
        
        print("✅ Metrics data structure created")
        print(f"Metrics data: {metrics_data}")
        
        # Test 8: Test enumeration like in Word export 
        print("8. Testing enumeration...")
        for i, (metric, value) in enumerate(metrics_data, 1):
            print(f"Row {i}: {metric} = {value}")
        print("✅ Enumeration test passed")
        
        print("\n🎉 All tests passed! Word export should work.")
        
except Exception as e:
    print(f"❌ Error during test: {e}")
    import traceback
    traceback.print_exc()
