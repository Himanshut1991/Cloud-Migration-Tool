#!/usr/bin/env python3
"""Minimal working backend with inline export functions"""

import os
import json
from datetime import datetime
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy

# Simple working Flask app
app = Flask(__name__)
CORS(app)

# Database setup
basedir = os.path.dirname(os.path.abspath(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(basedir, "migration_tool.db")}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# Initialize models
try:
    from models_new import init_models
    models = init_models(db)
    print("✅ Models initialized")
except Exception as e:
    print(f"❌ Models error: {e}")
    models = {}

# Create exports directory
exports_dir = os.path.join(basedir, 'exports')
os.makedirs(exports_dir, exist_ok=True)

@app.route('/api/dashboard', methods=['GET'])
def dashboard():
    """Dashboard endpoint"""
    try:
        # Get basic counts
        Server = models.get('Server')
        Database = models.get('Database')
        FileShare = models.get('FileShare')
        
        servers_count = Server.query.count() if Server else 0
        databases_count = Database.query.count() if Database else 0
        file_shares_count = FileShare.query.count() if FileShare else 0
        
        return jsonify({
            'infrastructure_summary': {
                'servers': servers_count,
                'databases': databases_count,
                'file_shares': file_shares_count,
                'total_components': servers_count + databases_count + file_shares_count
            },
            'cost_breakdown': [
                {'category': 'Compute', 'cost': servers_count * 200},
                {'category': 'Storage', 'cost': file_shares_count * 100},
                {'category': 'Database', 'cost': databases_count * 300},
                {'category': 'Network', 'cost': 500}
            ],
            'migration_timeline': {
                'total_weeks': max(8, servers_count + databases_count + file_shares_count),
                'phases': ['Planning', 'Setup', 'Migration', 'Testing', 'Go-Live']
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/export', methods=['POST'])
def export_report():
    """Export with inline implementations"""
    try:
        data = request.get_json() or {}
        format_type = data.get('format', 'pdf').lower()
        
        print(f"📄 Processing {format_type} export request...")
        
        # Get basic data
        Server = models.get('Server')
        Database = models.get('Database')
        FileShare = models.get('FileShare')
        
        servers_count = Server.query.count() if Server else 0
        databases_count = Database.query.count() if Database else 0
        file_shares_count = FileShare.query.count() if FileShare else 0
        total_components = servers_count + databases_count + file_shares_count
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'migration_plan_{timestamp}.{format_type}'
        filepath = os.path.join(exports_dir, filename)
        
        if format_type == 'pdf':
            # Create PDF
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib import colors
            
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Content
            story.append(Paragraph("Cloud Migration Assessment Report", styles['Title']))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            story.append(Paragraph("Infrastructure Summary", styles['Heading1']))
            story.append(Paragraph(f"Total Components: {total_components}", styles['Normal']))
            story.append(Paragraph(f"Servers: {servers_count}, Databases: {databases_count}, File Shares: {file_shares_count}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Simple table
            data = [
                ['Component', 'Count', 'Priority'],
                ['Servers', str(servers_count), 'High'],
                ['Databases', str(databases_count), 'Critical'],
                ['File Shares', str(file_shares_count), 'Medium']
            ]
            
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(table)
            
            doc.build(story)
            
        elif format_type == 'excel':
            # Create Excel
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Migration Report"
            
            # Headers
            ws['A1'] = "Cloud Migration Assessment Report"
            ws['A1'].font = Font(bold=True, size=16)
            ws['A2'] = f"Generated: {datetime.now().strftime('%B %d, %Y')}"
            
            # Data
            ws['A4'] = "Infrastructure Summary"
            ws['A4'].font = Font(bold=True, size=14)
            
            ws['A6'] = "Component"
            ws['B6'] = "Count"
            ws['C6'] = "Priority"
            
            for col in ['A6', 'B6', 'C6']:
                ws[col].font = Font(bold=True)
                ws[col].fill = PatternFill(start_color="D6EAF8", end_color="D6EAF8", fill_type="solid")
            
            ws['A7'] = "Servers"
            ws['B7'] = servers_count
            ws['C7'] = "High"
            
            ws['A8'] = "Databases"
            ws['B8'] = databases_count
            ws['C8'] = "Critical"
            
            ws['A9'] = "File Shares"
            ws['B9'] = file_shares_count
            ws['C9'] = "Medium"
            
            # Auto-adjust column width
            for col in ['A', 'B', 'C']:
                ws.column_dimensions[col].width = 20
                
            wb.save(filepath)
            
        elif format_type == 'word':
            # Create Word
            from docx import Document
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Cloud Migration Assessment Report', 0)
            title.alignment = 1
            
            # Date
            date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = 1
            
            # Summary
            doc.add_heading('Infrastructure Summary', level=1)
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"Total Components: {total_components}")
            summary_para.add_run(f"\nServers: {servers_count}")
            summary_para.add_run(f"\nDatabases: {databases_count}")
            summary_para.add_run(f"\nFile Shares: {file_shares_count}")
            
            # Table
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Table Grid'
            
            # Headers
            table.rows[0].cells[0].text = 'Component'
            table.rows[0].cells[1].text = 'Count'
            table.rows[0].cells[2].text = 'Priority'
            
            # Data
            table.rows[1].cells[0].text = 'Servers'
            table.rows[1].cells[1].text = str(servers_count)
            table.rows[1].cells[2].text = 'High'
            
            table.rows[2].cells[0].text = 'Databases'
            table.rows[2].cells[1].text = str(databases_count)
            table.rows[2].cells[2].text = 'Critical'
            
            table.rows[3].cells[0].text = 'File Shares'
            table.rows[3].cells[1].text = str(file_shares_count)
            table.rows[3].cells[2].text = 'Medium'
            
            doc.save(filepath)
            
        else:
            return jsonify({'error': f'Unsupported format: {format_type}'}), 400
        
        # Return success
        file_size = os.path.getsize(filepath)
        
        print(f"✅ {format_type.upper()} export completed: {filename}")
        
        return jsonify({
            'message': f'{format_type.upper()} export completed successfully',
            'format': format_type,
            'filename': filename,
            'filepath': filepath,
            'file_size': file_size,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"❌ Export failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Export failed: {str(e)}'}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download exported file"""
    try:
        filepath = os.path.join(exports_dir, filename)
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    with app.app_context():
        try:
            db.create_all()
            print("✅ Database initialized")
        except Exception as e:
            print(f"⚠️ Database warning: {e}")
    
    print("🚀 Starting minimal backend server...")
    print("📍 Server: http://localhost:5000")
    print("📄 Endpoints: /api/dashboard, /api/export, /api/download/<filename>")
    app.run(host='127.0.0.1', port=5000, debug=True)
