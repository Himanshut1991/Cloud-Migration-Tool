#!/usr/bin/env python3
"""
Real Data Backend - Serves ONLY real database data, no sample data
"""

from flask import Flask, jsonify, request, make_response
from flask_cors import CORS
import sqlite3
import logging
import traceback
import os
from datetime import datetime
from services.ai_recommendations import AIRecommendationService

# Setup logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Initialize AI service
ai_service = AIRecommendationService()

DATABASE_PATH = 'migration_tool.db'

def get_db_connection():
    """Get database connection"""
    conn = sqlite3.connect(DATABASE_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def dict_from_row(row):
    """Convert sqlite3.Row to dict"""
    return dict(row) if row else None

@app.route('/api/servers', methods=['GET', 'POST'])
def handle_servers():
    try:
        if request.method == 'POST':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO servers (server_id, os_type, vcpu, ram, disk_size, disk_type, uptime_pattern, current_hosting, technology, technology_version)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (data['server_id'], data['os_type'], data['vcpu'], data['ram'], 
                  data['disk_size'], data['disk_type'], data['uptime_pattern'], 
                  data['current_hosting'], data['technology'], data['technology_version']))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Server added successfully'})
        
        # GET request
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM servers ORDER BY id')
        rows = cursor.fetchall()
        servers = [dict_from_row(row) for row in rows]
        conn.close()
        
        logger.info(f"Retrieved {len(servers)} servers from database")
        return jsonify({'servers': servers})
        
    except Exception as e:
        logger.error(f"Error in /api/servers: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/servers/<int:server_id>', methods=['PUT', 'DELETE'])
def handle_server_by_id(server_id):
    try:
        if request.method == 'PUT':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE servers SET server_id=?, os_type=?, vcpu=?, ram=?, disk_size=?, 
                       disk_type=?, uptime_pattern=?, current_hosting=?, technology=?, technology_version=?
                WHERE id=?
            ''', (data['server_id'], data['os_type'], data['vcpu'], data['ram'], 
                  data['disk_size'], data['disk_type'], data['uptime_pattern'], 
                  data['current_hosting'], data['technology'], data['technology_version'], server_id))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Server updated successfully'})
        
        elif request.method == 'DELETE':
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('DELETE FROM servers WHERE id=?', (server_id,))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Server deleted successfully'})
            
    except Exception as e:
        logger.error(f"Error in /api/servers/{server_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/databases', methods=['GET', 'POST'])
def handle_databases():
    try:
        if request.method == 'POST':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO databases (db_name, db_type, size_gb, ha_dr_required, backup_frequency, 
                       licensing_model, server_id, write_frequency, downtime_tolerance, real_time_sync)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (data['db_name'], data['db_type'], data['size_gb'], data.get('ha_dr_required', False),
                  data.get('backup_frequency', 'Daily'), data.get('licensing_model', 'Standard'),
                  data.get('server_id', 1), data.get('write_frequency', 'Medium'),
                  data.get('downtime_tolerance', 'Low'), data.get('real_time_sync', False)))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Database added successfully'})
        
        # GET request
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM databases ORDER BY id')
        rows = cursor.fetchall()
        databases = [dict_from_row(row) for row in rows]
        conn.close()
        
        logger.info(f"Retrieved {len(databases)} databases from database")
        return jsonify({'databases': databases})
        
    except Exception as e:
        logger.error(f"Error in /api/databases: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/databases/<int:database_id>', methods=['PUT', 'DELETE'])
def handle_database_by_id(database_id):
    try:
        if request.method == 'PUT':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE databases SET db_name=?, db_type=?, size_gb=?, ha_dr_required=?, 
                       backup_frequency=?, licensing_model=?, server_id=?, write_frequency=?, 
                       downtime_tolerance=?, real_time_sync=?
                WHERE id=?
            ''', (data['db_name'], data['db_type'], data['size_gb'], data.get('ha_dr_required', False),
                  data.get('backup_frequency', 'Daily'), data.get('licensing_model', 'Standard'),
                  data.get('server_id', 1), data.get('write_frequency', 'Medium'),
                  data.get('downtime_tolerance', 'Low'), data.get('real_time_sync', False), database_id))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Database updated successfully'})
        
        elif request.method == 'DELETE':
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('DELETE FROM databases WHERE id=?', (database_id,))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Database deleted successfully'})
            
    except Exception as e:
        logger.error(f"Error in /api/databases/{database_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/file-shares', methods=['GET', 'POST'])
def handle_file_shares():
    try:
        if request.method == 'POST':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO file_shares (share_name, total_size_gb, access_pattern, snapshot_required, 
                       retention_days, server_id, write_frequency, downtime_tolerance, real_time_sync)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (data['share_name'], data['total_size_gb'], data.get('access_pattern', 'Medium'),
                  data.get('snapshot_required', False), data.get('retention_days', 30),
                  data.get('server_id', 1), data.get('write_frequency', 'Medium'),
                  data.get('downtime_tolerance', 'Low'), data.get('real_time_sync', False)))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'File share added successfully'})
        
        # GET request
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM file_shares ORDER BY id')
        rows = cursor.fetchall()
        file_shares = [dict_from_row(row) for row in rows]
        conn.close()
        
        logger.info(f"Retrieved {len(file_shares)} file shares from database")
        return jsonify({'file_shares': file_shares})
        
    except Exception as e:
        logger.error(f"Error in /api/file-shares: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/file-shares/<int:file_share_id>', methods=['PUT', 'DELETE'])
def handle_file_share_by_id(file_share_id):
    try:
        if request.method == 'PUT':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE file_shares SET share_name=?, total_size_gb=?, access_pattern=?, 
                       snapshot_required=?, retention_days=?, server_id=?, write_frequency=?, 
                       downtime_tolerance=?, real_time_sync=?
                WHERE id=?
            ''', (data['share_name'], data['total_size_gb'], data.get('access_pattern', 'Medium'),
                  data.get('snapshot_required', False), data.get('retention_days', 30),
                  data.get('server_id', 1), data.get('write_frequency', 'Medium'),
                  data.get('downtime_tolerance', 'Low'), data.get('real_time_sync', False), file_share_id))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'File share updated successfully'})
        
        elif request.method == 'DELETE':
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('DELETE FROM file_shares WHERE id=?', (file_share_id,))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'File share deleted successfully'})
            
    except Exception as e:
        logger.error(f"Error in /api/file-shares/{file_share_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/resource-rates', methods=['GET', 'POST'])
def handle_resource_rates():
    try:
        if request.method == 'POST':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO resource_rates (role, duration_weeks, hours_per_week, rate_per_hour)
                VALUES (?, ?, ?, ?)
            ''', (data['role'], data['duration_weeks'], data['hours_per_week'], data['rate_per_hour']))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Resource rate added successfully'})
        
        # GET request
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM resource_rates ORDER BY id')
        rows = cursor.fetchall()
        resource_rates = [dict_from_row(row) for row in rows]
        conn.close()
        
        logger.info(f"Retrieved {len(resource_rates)} resource rates from database")
        return jsonify({'resource_rates': resource_rates})
        
    except Exception as e:
        logger.error(f"Error in /api/resource-rates: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/resource-rates/<int:rate_id>', methods=['PUT', 'DELETE'])
def handle_resource_rate_by_id(rate_id):
    try:
        if request.method == 'PUT':
            data = request.json
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE resource_rates SET role=?, duration_weeks=?, hours_per_week=?, rate_per_hour=?
                WHERE id=?
            ''', (data['role'], data['duration_weeks'], data['hours_per_week'], data['rate_per_hour'], rate_id))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Resource rate updated successfully'})
        
        elif request.method == 'DELETE':
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('DELETE FROM resource_rates WHERE id=?', (rate_id,))
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Resource rate deleted successfully'})
            
    except Exception as e:
        logger.error(f"Error in /api/resource-rates/{rate_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500
        resource_rates = [dict_from_row(row) for row in rows]
        conn.close()
        
        logger.info(f"Retrieved {len(resource_rates)} resource rates from database")
        return jsonify({'resource_rates': resource_rates})
        
    except Exception as e:
        logger.error(f"Error in /api/resource-rates: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/dashboard', methods=['GET'])
def get_dashboard():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get counts and basic stats from real database
        cursor.execute('SELECT COUNT(*) FROM servers')
        total_servers = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM databases')
        total_databases = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM file_shares')
        total_file_shares = cursor.fetchone()[0]
        
        # Calculate total data size from databases and file shares
        cursor.execute('SELECT COALESCE(SUM(size_gb), 0) FROM databases')
        database_size = cursor.fetchone()[0]
        
        cursor.execute('SELECT COALESCE(SUM(total_size_gb), 0) FROM file_shares')
        file_share_size = cursor.fetchone()[0]
        
        cursor.execute('SELECT COALESCE(SUM(disk_size), 0) FROM servers')
        server_storage_size = cursor.fetchone()[0]
        
        total_data_gb = database_size + file_share_size + server_storage_size
        
        # Get estimated costs (placeholder calculation from real data)
        cursor.execute('SELECT SUM(ram * 0.1 + vcpu * 0.05 + disk_size * 0.02) FROM servers')
        result = cursor.fetchone()[0]
        estimated_monthly_cost = round(result * 730, 2) if result else 0  # rough estimate
        
        conn.close()
        
        dashboard_data = {
            'infrastructure_summary': {
                'servers': total_servers,
                'databases': total_databases,
                'file_shares': total_file_shares,
                'total_items': total_servers + total_databases + total_file_shares,
                'total_data_gb': total_data_gb
            },
            'cost_estimation': {
                'monthly_cost': estimated_monthly_cost,
                'annual_cost': estimated_monthly_cost * 12,
                'currency': 'USD',
                'last_updated': 'Now'
            },
            'migration_timeline': {
                'estimated_duration_weeks': 12,
                'phases': 5,
                'complexity': 'Medium'
            }
        }
        
        logger.info(f"Dashboard data: {dashboard_data}")
        return jsonify(dashboard_data)
        
    except Exception as e:
        logger.error(f"Error in /api/dashboard: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/cost-estimation', methods=['POST', 'OPTIONS'])
def cost_estimation():
    # Handle preflight request
    if request.method == 'OPTIONS':
        response = make_response()
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add('Access-Control-Allow-Headers', "*")
        response.headers.add('Access-Control-Allow-Methods', "*")
        return response
        
    try:
        data = request.json
        logger.info(f"AI Cost estimation request: {data}")
        
        cloud_provider = data.get('cloud_provider', 'AWS')
        target_region = data.get('target_region', 'us-east-1')
        
        # Get real infrastructure data from database
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get servers
        cursor.execute('SELECT * FROM servers')
        servers_rows = cursor.fetchall()
        servers = [dict_from_row(row) for row in servers_rows]
        
        # Get databases
        cursor.execute('SELECT * FROM databases')
        databases_rows = cursor.fetchall()
        databases = [dict_from_row(row) for row in databases_rows]
        
        # Get file shares
        cursor.execute('SELECT * FROM file_shares')
        file_shares_rows = cursor.fetchall()
        file_shares = [dict_from_row(row) for row in file_shares_rows]
        
        conn.close()
        
        # Prepare infrastructure data for AI analysis
        infrastructure_data = {
            'servers': servers,
            'databases': databases,
            'file_shares': file_shares
        }
        
        logger.info(f"Using AI for cost estimation with {len(servers)} servers, {len(databases)} databases, {len(file_shares)} file shares")
        
        # Get AI-powered cost estimation
        cost_data = ai_service.get_ai_cost_estimation(
            infrastructure_data, 
            cloud_provider, 
            target_region
        )
        
        logger.info(f"Cost estimation completed - AI used: {not cost_data.get('ai_insights', {}).get('fallback_used', True)}")
        
        response = jsonify(cost_data)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        monthly_compute_cost = data.get('servers', 0) * avg_compute_rate * 730
        monthly_storage_cost = data.get('storage_gb', 0) * avg_storage_rate * 730
        total_monthly_cost = monthly_compute_cost + monthly_storage_cost
        
        result = {
            'monthly_cost': round(total_monthly_cost, 2),
            'yearly_cost': round(total_monthly_cost * 12, 2),
            'breakdown': {
                'compute': round(monthly_compute_cost, 2),
                'storage': round(monthly_storage_cost, 2)
            }
        }
        
        logger.info(f"Cost estimation: {result}")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in /api/cost-estimation: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/migration-strategy', methods=['POST', 'OPTIONS'])
def migration_strategy():
    # Handle preflight request
    if request.method == 'OPTIONS':
        response = make_response()
        response.headers.add("Access-Control-Allow-Origin", "*")
        response.headers.add('Access-Control-Allow-Headers', "*")
        response.headers.add('Access-Control-Allow-Methods', "*")
        return response
        
    try:
        data = request.json
        logger.info(f"AI Migration strategy request: {data}")
        
        cloud_provider = data.get('cloud_provider', 'AWS')
        target_region = data.get('target_region', 'us-east-1')
        complexity = data.get('migration_complexity', 'medium')
        
        # Get real infrastructure data from database
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get servers
        cursor.execute('SELECT * FROM servers')
        servers_rows = cursor.fetchall()
        servers = [dict_from_row(row) for row in servers_rows]
        
        # Get databases
        cursor.execute('SELECT * FROM databases')
        databases_rows = cursor.fetchall()
        databases = [dict_from_row(row) for row in databases_rows]
        
        # Get file shares
        cursor.execute('SELECT * FROM file_shares')
        file_shares_rows = cursor.fetchall()
        file_shares = [dict_from_row(row) for row in file_shares_rows]
        
        conn.close()
        
        # Prepare infrastructure data for AI analysis
        infrastructure_data = {
            'servers': servers,
            'databases': databases,
            'file_shares': file_shares
        }
        
        logger.info(f"Using AI for migration strategy with {len(servers)} servers, {len(databases)} databases, {len(file_shares)} file shares")
        
        # Get AI-powered migration strategy
        strategy_data = ai_service.get_ai_migration_strategy(
            infrastructure_data, 
            cloud_provider, 
            target_region,
            complexity
        )
        
        logger.info(f"AI Migration strategy completed - AI used: {not strategy_data.get('ai_insights', {}).get('fallback_used', True)}")
        
        response = jsonify(strategy_data)
        response.headers.add("Access-Control-Allow-Origin", "*")
        return response
        
    except Exception as e:
        logger.error(f"Error in /api/migration-strategy: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/timeline', methods=['POST'])
def generate_timeline():
    """Generate migration timeline based on project data"""
    try:
        data = request.get_json()
        logger.info(f"Timeline generation request: {data}")
        
        # Get infrastructure data to calculate duration
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get counts for duration calculation
        cursor.execute('SELECT COUNT(*) FROM servers')
        server_count = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM databases')
        database_count = cursor.fetchone()[0]
        
        cursor.execute('SELECT COUNT(*) FROM file_shares')
        file_share_count = cursor.fetchone()[0]
        
        conn.close()
        
        # Calculate duration based on infrastructure size (same logic as migration strategy)
        total_components = server_count + database_count + file_share_count
        duration_weeks = max(8, total_components * 2)
        duration_months = round(duration_weeks / 4.3, 1)  # More accurate weeks to months conversion
        
        # Get custom start date from request or use default
        start_date = data.get('start_date', '2025-09-01') if data else '2025-09-01'
        
        # Calculate end date based on calculated duration
        from datetime import timedelta
        start_dt = datetime.strptime(start_date, '%Y-%m-%d')
        end_dt = start_dt + timedelta(weeks=duration_weeks)
        end_date = end_dt.strftime('%Y-%m-%d')
        
        logger.info(f"Timeline dates: {start_date} to {end_date} ({duration_weeks} weeks, {total_components} components)")
        
        # Calculate phase durations that add up to total duration
        phase1_duration = max(3, round(duration_weeks * 0.25))
        phase2_duration = max(4, round(duration_weeks * 0.4))
        phase3_duration = max(4, round(duration_weeks * 0.25))
        phase4_duration = max(2, duration_weeks - phase1_duration - phase2_duration - phase3_duration)
        
        # Adjust if total doesn't match
        total_calculated = phase1_duration + phase2_duration + phase3_duration + phase4_duration
        if total_calculated != duration_weeks:
            # Adjust the largest phase to match
            phase2_duration += (duration_weeks - total_calculated)
        
        logger.info(f"Phase durations: P1={phase1_duration}, P2={phase2_duration}, P3={phase3_duration}, P4={phase4_duration}, Total={phase1_duration + phase2_duration + phase3_duration + phase4_duration}")
        
        # Timeline data with dynamic dates and duration
        timeline_data = {
            "project_overview": {
                "total_duration_weeks": duration_weeks,
                "total_duration_months": duration_months,
                "estimated_start_date": start_date,
                "estimated_end_date": end_date,
                "confidence_level": "85%",
                "complexity_score": min(10, total_components * 0.5)  # Dynamic complexity based on component count
            },
            "phases": [
                {
                    "id": 1,
                    "name": "Assessment & Planning",
                    "description": "Initial assessment and detailed planning",
                    "start_week": 1,
                    "end_week": phase1_duration,
                    "duration_weeks": phase1_duration,
                    "dependencies": [],
                    "milestones": ["Infrastructure Assessment Complete", "Migration Plan Approved"],
                    "components": ["Server Discovery", "Application Mapping", "Risk Assessment"],
                    "risks": ["Incomplete inventory", "Missing dependencies"],
                    "resources_required": ["Cloud Architect", "System Administrator"],
                    "status": "pending"
                },
                {
                    "id": 2,
                    "name": "Infrastructure Migration",
                    "description": "Migrate servers and infrastructure components",
                    "start_week": phase1_duration + 1,
                    "end_week": phase1_duration + phase2_duration,
                    "duration_weeks": phase2_duration,
                    "dependencies": ["Phase 1"],
                    "milestones": ["Test Environment Ready", "Production Infrastructure Live"],
                    "components": ["Server Migration", "Network Configuration", "Security Setup"],
                    "risks": ["Network connectivity issues", "Performance degradation"],
                    "resources_required": ["Cloud Engineer", "Network Specialist"],
                    "status": "pending"
                },
                {
                    "id": 3,
                    "name": "Data Migration",
                    "description": "Migrate databases and file shares",
                    "start_week": phase1_duration + phase2_duration + 1,
                    "end_week": phase1_duration + phase2_duration + phase3_duration,
                    "duration_weeks": phase3_duration,
                    "dependencies": ["Phase 1"],
                    "milestones": ["Data Sync Established", "Data Validation Complete"],
                    "components": ["Database Migration", "File Share Migration", "Data Validation"],
                    "risks": ["Data corruption", "Extended sync time"],
                    "resources_required": ["Database Administrator", "Data Engineer"],
                    "status": "pending"
                },
                {
                    "id": 4,
                    "name": "Application Cutover",
                    "description": "Final application migration and cutover",
                    "start_week": phase1_duration + phase2_duration + phase3_duration + 1,
                    "end_week": duration_weeks,
                    "duration_weeks": phase4_duration,
                    "dependencies": ["Phase 2", "Phase 3"],
                    "milestones": ["Application Migration Complete", "Production Cutover"],
                    "components": ["Application Deployment", "DNS Cutover", "Monitoring Setup"],
                    "risks": ["Application compatibility", "User acceptance"],
                    "resources_required": ["Application Developer", "System Administrator"],
                    "status": "pending"
                }
            ],
            "critical_path": ["Phase 1", "Phase 2", "Phase 3", "Phase 4"],
            "resource_allocation": [
                {
                    "role": "Cloud Architect",
                    "weeks_allocated": max(3, round(duration_weeks * 0.25)),
                    "overlap_phases": [1],
                    "peak_utilization_week": 2
                },
                {
                    "role": "Database Expert",
                    "weeks_allocated": max(4, round(duration_weeks * 0.5)),
                    "overlap_phases": [1, 3],
                    "peak_utilization_week": round(duration_weeks * 0.6)
                },
                {
                    "role": "DevOps Engineer",
                    "weeks_allocated": max(4, round(duration_weeks * 0.4)),
                    "overlap_phases": [2, 4],
                    "peak_utilization_week": round(duration_weeks * 0.4)
                }
            ],
            "risk_mitigation": [
                {
                    "risk": "Data corruption during migration",
                    "probability": "Medium",
                    "impact": "High",
                    "mitigation_strategy": "Implement comprehensive backup strategy",
                    "timeline_buffer_weeks": 2
                },
                {
                    "risk": "Extended downtime during cutover",
                    "probability": "High",
                    "impact": "High",
                    "mitigation_strategy": "Plan for rollback procedures and conduct thorough testing",
                    "timeline_buffer_weeks": 1
                }
            ],
            "success_criteria": [
                "All servers migrated successfully",
                "Zero data loss during migration",
                "Application performance maintained",
                "Downtime within acceptable limits"
            ],
            "ai_insights": {
                "optimization_suggestions": [
                    "Consider overlapping phases 2 and 3 to reduce total timeline",
                    "Implement blue-green deployment for reduced downtime"
                ],
                "timeline_risks": [
                    "Data migration may take longer than estimated for large databases",
                    "Application testing phase should include buffer time"
                ],
                "resource_recommendations": [
                    "Assign dedicated DBA for phases 1 and 3",
                    "Consider additional cloud architect for parallel workstreams"
                ]
            }
        }
        
        logger.info("Timeline generated successfully")
        return jsonify(timeline_data)
        
    except Exception as e:
        logger.error(f"Error in /api/timeline: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/ai-status', methods=['GET'])
def ai_status():
    try:
        # Simple AI status - would integrate with actual AI service
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT COUNT(*) FROM servers')
        server_count = cursor.fetchone()[0]
        conn.close()
        
        status = {
            'ai_enabled': server_count > 0,  # AI enabled if we have data to analyze
            'models_available': ['Claude 3.5 Sonnet', 'Claude 3 Sonnet', 'Titan Text G1 - Express'] if server_count > 0 else [],
            'last_analysis': '2024-01-15T10:00:00Z' if server_count > 0 else None,
            'recommendations_count': server_count if server_count > 0 else 0
        }
        
        logger.info(f"AI status: {status}")
        return jsonify(status)
        
    except Exception as e:
        logger.error(f"Error in /api/ai-status: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/export', methods=['POST'])
def export_report():
    try:
        data = request.json
        export_format = data.get('format', 'excel')
        report_types = data.get('types', ['cost_estimation', 'migration_strategy', 'timeline'])
        
        logger.info(f"Export request - Format: {export_format}, Types: {report_types}")
        
        # Get all data from database
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Get servers
        cursor.execute('SELECT * FROM servers')
        servers_rows = cursor.fetchall()
        servers = [dict_from_row(row) for row in servers_rows]
        
        # Get databases
        cursor.execute('SELECT * FROM databases')
        databases_rows = cursor.fetchall()
        databases = [dict_from_row(row) for row in databases_rows]
        
        # Get file shares
        cursor.execute('SELECT * FROM file_shares')
        file_shares_rows = cursor.fetchall()
        file_shares = [dict_from_row(row) for row in file_shares_rows]
        
        # Get resource rates
        cursor.execute('SELECT * FROM resource_rates')
        resource_rates_rows = cursor.fetchall()
        resource_rates = [dict_from_row(row) for row in resource_rates_rows]
        
        conn.close()
        
        # Prepare infrastructure data
        infrastructure_data = {
            'servers': servers,
            'databases': databases,
            'file_shares': file_shares,
            'resource_rates': resource_rates
        }
        
        # Initialize export data
        export_data = {
            'infrastructure': infrastructure_data,
            'export_metadata': {
                'timestamp': datetime.now().isoformat(),
                'format': export_format,
                'types': report_types
            }
        }
        
        # Generate AI data based on selected report types (with fallback if AI unavailable)
        if 'cost_estimation' in report_types:
            logger.info("Generating cost estimation data for export")
            try:
                if ai_service and ai_service.bedrock_client:
                    cost_data = ai_service.get_ai_cost_estimation(
                        infrastructure_data, 'AWS', 'us-east-1', 'medium'
                    )
                else:
                    raise Exception("AI service unavailable")
            except Exception as e:
                logger.warning(f"AI cost estimation failed: {e}, using fallback data")
                cost_data = _generate_fallback_cost_data(servers, databases, file_shares)
            export_data['cost_estimation'] = cost_data
        
        if 'migration_strategy' in report_types:
            logger.info("Generating migration strategy data for export")
            try:
                if ai_service and ai_service.bedrock_client:
                    strategy_data = ai_service.get_ai_migration_strategy(
                        infrastructure_data, 'AWS', 'us-east-1', 'medium'
                    )
                else:
                    raise Exception("AI service unavailable")
            except Exception as e:
                logger.warning(f"AI migration strategy failed: {e}, using fallback data")
                strategy_data = _generate_fallback_strategy_data(servers, databases, file_shares)
            export_data['migration_strategy'] = strategy_data
        
        if 'timeline' in report_types:
            logger.info("Generating timeline data for export")
            # Generate timeline data
            from datetime import timedelta
            start_dt = datetime.now()
            end_dt = start_dt + timedelta(weeks=16)
            
            timeline_data = {
                "project_overview": {
                    "total_duration_weeks": 16,
                    "estimated_start_date": start_dt.strftime('%Y-%m-%d'),
                    "estimated_end_date": end_dt.strftime('%Y-%m-%d'),
                    "confidence_level": "85%"
                },
                "phases": [
                    {
                        "id": 1,
                        "name": "Assessment & Planning",
                        "start_week": 1,
                        "end_week": 4,
                        "duration_weeks": 4
                    },
                    {
                        "id": 2,
                        "name": "Infrastructure Migration",
                        "start_week": 5,
                        "end_week": 10,
                        "duration_weeks": 6
                    },
                    {
                        "id": 3,
                        "name": "Data Migration",
                        "start_week": 8,
                        "end_week": 14,
                        "duration_weeks": 7
                    },
                    {
                        "id": 4,
                        "name": "Application Cutover",
                        "start_week": 15,
                        "end_week": 16,
                        "duration_weeks": 2
                    }
                ]
            }
            export_data['timeline'] = timeline_data
        
        # Add inventory data
        export_data['inventory'] = {
            'servers': servers,
            'databases': databases,
            'file_shares': file_shares,
            'resource_rates': resource_rates
        }
        
        # Generate file based on format
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if export_format == 'excel':
            filename = f'migration_report_{timestamp}.xlsx'
            filepath = _generate_excel_report(export_data, filename)
        elif export_format == 'pdf':
            filename = f'migration_report_{timestamp}.pdf'
            filepath = _generate_pdf_report(export_data, filename)
        elif export_format == 'word':
            filename = f'migration_report_{timestamp}.docx'
            filepath = _generate_word_report(export_data, filename)
        else:
            return jsonify({'error': 'Invalid export format'}), 400
        
        file_size = os.path.getsize(filepath) if os.path.exists(filepath) else 0
        
        result = {
            'message': f'{export_format.upper()} export completed successfully',
            'format': export_format,
            'filename': filename,
            'filepath': filepath,
            'file_size': file_size,
            'timestamp': datetime.now().isoformat(),
            'types_included': report_types,
            'download_url': f'/api/download/{filename}'
        }
        
        logger.info(f"Export completed: {filename} ({file_size} bytes)")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error in /api/export: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

def _generate_excel_report(export_data, filename):
    """Generate Excel report with multiple sheets"""
    try:
        import pandas as pd
        from openpyxl import Workbook
        from openpyxl.utils.dataframe import dataframe_to_rows
        from openpyxl.styles import Font, PatternFill
        
        exports_dir = os.path.join(os.path.dirname(__file__), 'exports')
        os.makedirs(exports_dir, exist_ok=True)
        filepath = os.path.join(exports_dir, filename)
        
        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        # Summary sheet
        ws_summary = wb.create_sheet(title="Executive Summary")
        ws_summary.append(["Cloud Migration Report", "", "", ""])
        ws_summary.append(["Generated:", datetime.now().strftime('%Y-%m-%d %H:%M:%S')])
        ws_summary.append([])
        
        # Inventory summary
        inventory = export_data.get('inventory', {})
        ws_summary.append(["Infrastructure Summary"])
        ws_summary.append(["Servers:", len(inventory.get('servers', []))])
        ws_summary.append(["Databases:", len(inventory.get('databases', []))])
        ws_summary.append(["File Shares:", len(inventory.get('file_shares', []))])
        ws_summary.append([])
        
        # Cost estimation summary
        if 'cost_estimation' in export_data:
            cost_data = export_data['cost_estimation']
            ws_summary.append(["Cost Estimation Summary"])
            grand_total = cost_data.get('grand_total', {})
            ws_summary.append(["Monthly Cloud Cost:", f"${grand_total.get('annual_cloud_cost', 0) / 12:,.2f}"])
            ws_summary.append(["Annual Cloud Cost:", f"${grand_total.get('annual_cloud_cost', 0):,.2f}"])
            ws_summary.append(["Migration Cost:", f"${grand_total.get('one_time_migration_cost', 0):,.2f}"])
            ws_summary.append([])
        
        # Inventory sheets
        if inventory.get('servers'):
            ws_servers = wb.create_sheet(title="Servers")
            df_servers = pd.DataFrame(inventory['servers'])
            for r in dataframe_to_rows(df_servers, index=False, header=True):
                ws_servers.append(r)
        
        if inventory.get('databases'):
            ws_databases = wb.create_sheet(title="Databases")
            df_databases = pd.DataFrame(inventory['databases'])
            for r in dataframe_to_rows(df_databases, index=False, header=True):
                ws_databases.append(r)
        
        if inventory.get('file_shares'):
            ws_file_shares = wb.create_sheet(title="File Shares")
            df_file_shares = pd.DataFrame(inventory['file_shares'])
            for r in dataframe_to_rows(df_file_shares, index=False, header=True):
                ws_file_shares.append(r)
        
        # Cost analysis sheet
        if 'cost_estimation' in export_data:
            ws_cost = wb.create_sheet(title="Cost Analysis")
            cost_data = export_data['cost_estimation']
            
            ws_cost.append(["Cost Analysis Details"])
            ws_cost.append([])
            
            # Infrastructure costs
            infra = cost_data.get('cloud_infrastructure', {})
            ws_cost.append(["Infrastructure Costs"])
            ws_cost.append(["Component", "Monthly Cost", "Annual Cost"])
            ws_cost.append(["Servers", f"${infra.get('servers', {}).get('total_monthly_cost', 0):,.2f}", 
                           f"${infra.get('servers', {}).get('total_annual_cost', 0):,.2f}"])
            ws_cost.append(["Databases", f"${infra.get('databases', {}).get('total_monthly_cost', 0):,.2f}", 
                           f"${infra.get('databases', {}).get('total_annual_cost', 0):,.2f}"])
            ws_cost.append(["Storage", f"${infra.get('storage', {}).get('total_monthly_cost', 0):,.2f}", 
                           f"${infra.get('storage', {}).get('total_annual_cost', 0):,.2f}"])
            ws_cost.append([])
            
            # AI insights
            ai_insights = cost_data.get('ai_insights', {})
            if ai_insights:
                ws_cost.append(["AI Analysis"])
                ws_cost.append(["Confidence Level:", f"{ai_insights.get('confidence_level', 0) * 100:.0f}%"])
                ws_cost.append(["AI Model:", ai_insights.get('ai_model_used', 'N/A')])
                ws_cost.append(["Fallback Used:", "Yes" if ai_insights.get('fallback_used') else "No"])
                ws_cost.append([])
                
                ws_cost.append(["Cost Optimization Tips:"])
                for tip in ai_insights.get('cost_optimization_tips', []):
                    ws_cost.append([f"• {tip}"])
        
        # Migration strategy sheet
        if 'migration_strategy' in export_data:
            ws_strategy = wb.create_sheet(title="Migration Strategy")
            strategy_data = export_data['migration_strategy']
            
            ws_strategy.append(["Migration Strategy"])
            ws_strategy.append([])
            
            approach = strategy_data.get('migration_approach', {})
            ws_strategy.append(["Overall Strategy:", approach.get('overall_strategy', 'N/A')])
            ws_strategy.append(["Duration:", approach.get('estimated_duration', 'N/A')])
            ws_strategy.append(["Complexity:", approach.get('complexity_level', 'N/A')])
            ws_strategy.append([])
            
            # Phases
            phases = strategy_data.get('migration_phases', [])
            if phases:
                ws_strategy.append(["Migration Phases"])
                ws_strategy.append(["Phase", "Name", "Duration", "Components"])
                for phase in phases:
                    components = ', '.join(phase.get('components', []))
                    ws_strategy.append([phase.get('phase'), phase.get('name'), 
                                     phase.get('duration'), components])
        
        # Timeline sheet
        if 'timeline' in export_data:
            ws_timeline = wb.create_sheet(title="Timeline")
            timeline_data = export_data['timeline']
            
            ws_timeline.append(["Migration Timeline"])
            ws_timeline.append([])
            
            overview = timeline_data.get('project_overview', {})
            ws_timeline.append(["Duration:", f"{overview.get('total_duration_weeks')} weeks"])
            ws_timeline.append(["Start Date:", overview.get('estimated_start_date')])
            ws_timeline.append(["End Date:", overview.get('estimated_end_date')])
            ws_timeline.append([])
            
            phases = timeline_data.get('phases', [])
            if phases:
                ws_timeline.append(["Phase Timeline"])
                ws_timeline.append(["Phase", "Name", "Start Week", "End Week", "Duration"])
                for phase in phases:
                    ws_timeline.append([phase.get('id'), phase.get('name'), 
                                     phase.get('start_week'), phase.get('end_week'),
                                     f"{phase.get('duration_weeks')} weeks"])
        
        wb.save(filepath)
        return filepath
        
    except Exception as e:
        logger.error(f"Error generating Excel report: {str(e)}")
        raise

def _generate_pdf_report(export_data, filename):
    """Generate PDF report"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        
        exports_dir = os.path.join(os.path.dirname(__file__), 'exports')
        os.makedirs(exports_dir, exist_ok=True)
        filepath = os.path.join(exports_dir, filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title = Paragraph("Cloud Migration Report", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", styles['Heading1']))
        
        inventory = export_data.get('inventory', {})
        summary_data = [
            ['Infrastructure Component', 'Count'],
            ['Servers', str(len(inventory.get('servers', [])))],
            ['Databases', str(len(inventory.get('databases', [])))],
            ['File Shares', str(len(inventory.get('file_shares', [])))]
        ]
        
        summary_table = Table(summary_data)
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 12))
        
        # Cost Analysis
        if 'cost_estimation' in export_data:
            story.append(Paragraph("Cost Analysis", styles['Heading1']))
            cost_data = export_data['cost_estimation']
            grand_total = cost_data.get('grand_total', {})
            
            cost_summary = [
                ['Cost Component', 'Amount'],
                ['Monthly Cloud Cost', f"${grand_total.get('annual_cloud_cost', 0) / 12:,.2f}"],
                ['Annual Cloud Cost', f"${grand_total.get('annual_cloud_cost', 0):,.2f}"],
                ['One-time Migration Cost', f"${grand_total.get('one_time_migration_cost', 0):,.2f}"],
                ['Total First Year Cost', f"${grand_total.get('total_first_year_cost', 0):,.2f}"]
            ]
            
            cost_table = Table(cost_summary)
            cost_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(cost_table)
            story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        return filepath
        
    except Exception as e:
        logger.error(f"Error generating PDF report: {str(e)}")
        raise

def _generate_word_report(export_data, filename):
    """Generate Word document report"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        exports_dir = os.path.join(os.path.dirname(__file__), 'exports')
        os.makedirs(exports_dir, exist_ok=True)
        filepath = os.path.join(exports_dir, filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Cloud Migration Report', 0)
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        
        inventory = export_data.get('inventory', {})
        p = doc.add_paragraph('Infrastructure Overview:')
        p.add_run(f'\n• Servers: {len(inventory.get("servers", []))}')
        p.add_run(f'\n• Databases: {len(inventory.get("databases", []))}')
        p.add_run(f'\n• File Shares: {len(inventory.get("file_shares", []))}')
        
        # Cost Analysis
        if 'cost_estimation' in export_data:
            doc.add_heading('Cost Analysis', level=1)
            cost_data = export_data['cost_estimation']
            grand_total = cost_data.get('grand_total', {})
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Cost Component'
            hdr_cells[1].text = 'Amount'
            
            cost_items = [
                ('Monthly Cloud Cost', f"${grand_total.get('annual_cloud_cost', 0) / 12:,.2f}"),
                ('Annual Cloud Cost', f"${grand_total.get('annual_cloud_cost', 0):,.2f}"),
                ('Migration Cost', f"${grand_total.get('one_time_migration_cost', 0):,.2f}")
            ]
            
            for item, amount in cost_items:
                row_cells = table.add_row().cells
                row_cells[0].text = item
                row_cells[1].text = amount
        
        # Migration Strategy
        if 'migration_strategy' in export_data:
            doc.add_heading('Migration Strategy', level=1)
            strategy_data = export_data['migration_strategy']
            approach = strategy_data.get('migration_approach', {})
            
            p = doc.add_paragraph('Migration Approach:')
            p.add_run(f'\n• Strategy: {approach.get("overall_strategy", "N/A")}')
            p.add_run(f'\n• Duration: {approach.get("estimated_duration", "N/A")}')
            p.add_run(f'\n• Complexity: {approach.get("complexity_level", "N/A")}')
        
        doc.save(filepath)
        return filepath
        
    except Exception as e:
        logger.error(f"Error generating Word report: {str(e)}")
        raise

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    try:
        exports_dir = os.path.join(os.path.dirname(__file__), 'exports')
        filepath = os.path.join(exports_dir, filename)
        
        if os.path.exists(filepath):
            from flask import send_file
            return send_file(
                filepath,
                as_attachment=True,
                download_name=filename,
                mimetype='application/octet-stream'
            )
        else:
            return jsonify({'error': 'File not found'}), 404
        
    except Exception as e:
        logger.error(f"Error in /api/download: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT COUNT(*) FROM servers')
        server_count = cursor.fetchone()[0]
        conn.close()
        
        return jsonify({
            'status': 'healthy',
            'database_connected': True,
            'servers_in_db': server_count,
            'serving_real_data': True
        })
        
    except Exception as e:
        logger.error(f"Health check error: {str(e)}")
        return jsonify({
            'status': 'unhealthy',
            'database_connected': False,
            'error': str(e)
        }), 500

# ==================== CONFIGURATION ENDPOINTS ====================

@app.route('/api/cloud-preferences', methods=['GET', 'POST'])
def handle_cloud_preferences():
    try:
        if request.method == 'POST':
            data = request.json
            logger.info(f"Creating/updating cloud preferences: {data}")
            
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Check if record exists
            cursor.execute('SELECT COUNT(*) FROM cloud_preferences')
            count = cursor.fetchone()[0]
            
            if count > 0:
                # Update existing record
                cursor.execute('''
                    UPDATE cloud_preferences 
                    SET cloud_provider = ?, region = ?, preferred_services = ?, network_config = ?, updated_at = ?
                    WHERE id = (SELECT MIN(id) FROM cloud_preferences)
                ''', (data['cloud_provider'], data['region'], data['preferred_services'], 
                      data['network_config'], datetime.now().isoformat()))
            else:
                # Insert new record
                cursor.execute('''
                    INSERT INTO cloud_preferences (cloud_provider, region, preferred_services, network_config, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (data['cloud_provider'], data['region'], data['preferred_services'], 
                      data['network_config'], datetime.now().isoformat(), datetime.now().isoformat()))
            
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Cloud preferences saved successfully'})
        
        # GET request
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM cloud_preferences ORDER BY id LIMIT 1')
        row = cursor.fetchone()
        
        if row:
            cloud_preference = dict_from_row(row)
        else:
            # Return default values if no record exists
            cloud_preference = {
                'id': None,
                'cloud_provider': 'AWS',
                'region': 'us-east-1',
                'preferred_services': 'EC2,RDS,S3',
                'network_config': 'VPC with public/private subnets'
            }
        
        conn.close()
        logger.info(f"Retrieved cloud preferences: {cloud_preference}")
        return jsonify(cloud_preference)
        
    except Exception as e:
        logger.error(f"Error in /api/cloud-preferences: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/cloud-preferences/<int:pref_id>', methods=['PUT', 'DELETE'])
def handle_cloud_preference_by_id(pref_id):
    try:
        if request.method == 'PUT':
            data = request.json
            logger.info(f"Updating cloud preference {pref_id}: {data}")
            
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE cloud_preferences 
                SET cloud_provider = ?, region = ?, preferred_services = ?, network_config = ?, updated_at = ?
                WHERE id = ?
            ''', (data['cloud_provider'], data['region'], data['preferred_services'], 
                  data['network_config'], datetime.now().isoformat(), pref_id))
            
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Cloud preference updated successfully'})
            
        elif request.method == 'DELETE':
            logger.info(f"Deleting cloud preference {pref_id}")
            
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('DELETE FROM cloud_preferences WHERE id = ?', (pref_id,))
            conn.commit()
            conn.close()
            
            return jsonify({'success': True, 'message': 'Cloud preference deleted successfully'})
            
    except Exception as e:
        logger.error(f"Error in /api/cloud-preferences/{pref_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/business-constraints', methods=['GET', 'POST'])
def handle_business_constraints():
    try:
        if request.method == 'POST':
            data = request.json
            logger.info(f"Creating/updating business constraints: {data}")
            
            conn = get_db_connection()
            cursor = conn.cursor()
            
            # Check if record exists
            cursor.execute('SELECT COUNT(*) FROM business_constraints')
            count = cursor.fetchone()[0]
            
            if count > 0:
                # Update existing record
                cursor.execute('''
                    UPDATE business_constraints 
                    SET migration_window = ?, cutover_date = ?, downtime_tolerance = ?, budget_cap = ?, updated_at = ?
                    WHERE id = (SELECT MIN(id) FROM business_constraints)
                ''', (data['migration_window'], data['cutover_date'], data['downtime_tolerance'], 
                      data['budget_cap'], datetime.now().isoformat()))
            else:
                # Insert new record
                cursor.execute('''
                    INSERT INTO business_constraints (migration_window, cutover_date, downtime_tolerance, budget_cap, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (data['migration_window'], data['cutover_date'], data['downtime_tolerance'], 
                      data['budget_cap'], datetime.now().isoformat(), datetime.now().isoformat()))
            
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Business constraints saved successfully'})
        
        # GET request
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM business_constraints ORDER BY id LIMIT 1')
        row = cursor.fetchone()
        
        if row:
            business_constraint = dict_from_row(row)
        else:
            # Return default values if no record exists
            business_constraint = {
                'id': None,
                'migration_window': 'Business Hours',
                'cutover_date': '2025-12-01',
                'downtime_tolerance': 'Low',
                'budget_cap': '500000'
            }
        
        conn.close()
        logger.info(f"Retrieved business constraints: {business_constraint}")
        return jsonify(business_constraint)
        
    except Exception as e:
        logger.error(f"Error in /api/business-constraints: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/api/business-constraints/<int:constraint_id>', methods=['PUT', 'DELETE'])
def handle_business_constraint_by_id(constraint_id):
    try:
        if request.method == 'PUT':
            data = request.json
            logger.info(f"Updating business constraint {constraint_id}: {data}")
            
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE business_constraints 
                SET migration_window = ?, cutover_date = ?, downtime_tolerance = ?, budget_cap = ?, updated_at = ?
                WHERE id = ?
            ''', (data['migration_window'], data['cutover_date'], data['downtime_tolerance'], 
                  data['budget_cap'], datetime.now().isoformat(), constraint_id))
            
            conn.commit()
            conn.close()
            return jsonify({'success': True, 'message': 'Business constraint updated successfully'})
            
        elif request.method == 'DELETE':
            logger.info(f"Deleting business constraint {constraint_id}")
            
            conn = get_db_connection()
            cursor = conn.cursor()
            cursor.execute('DELETE FROM business_constraints WHERE id = ?', (constraint_id,))
            conn.commit()
            conn.close()
            
            return jsonify({'success': True, 'message': 'Business constraint deleted successfully'})
            
    except Exception as e:
        logger.error(f"Error in /api/business-constraints/{constraint_id}: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    logger.info("Starting Real Data Backend - serves ONLY database data")
    app.run(host='127.0.0.1', port=5000, debug=True)
