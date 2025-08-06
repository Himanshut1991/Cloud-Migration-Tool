#!/usr/bin/env python3
"""
Simple Export Service - Minimal working version
"""
import os
from datetime import datetime
import json

class SimpleExportService:
    """Simple export service with minimal functionality"""
    
    def __init__(self):
        self.output_dir = os.path.join(os.path.dirname(__file__), '..', 'exports')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_to_excel(self):
        """Export basic data to Excel"""
        try:
            import pandas as pd
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.xlsx'
            filepath = os.path.join(self.output_dir, filename)
            
            # Simple sample data
            data = {
                'Servers': [
                    {'Name': 'web-server-01', 'CPU': 4, 'RAM_GB': 16, 'Storage_GB': 100},
                    {'Name': 'db-server-01', 'CPU': 8, 'RAM_GB': 32, 'Storage_GB': 500},
                    {'Name': 'app-server-01', 'CPU': 6, 'RAM_GB': 24, 'Storage_GB': 200},
                ],
                'Summary': [
                    {'Total_Servers': 3, 'Total_RAM_GB': 72, 'Total_Storage_GB': 800},
                    {'Estimated_Monthly_Cost': '$2,450', 'Migration_Timeline': '3 months'},
                ]
            }
            
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Create servers sheet
                servers_df = pd.DataFrame(data['Servers'])
                servers_df.to_excel(writer, sheet_name='Servers', index=False)
                
                # Create summary sheet
                summary_df = pd.DataFrame(data['Summary'])
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            return filename
            
        except Exception as e:
            raise Exception(f"Excel export failed: {str(e)}")
    
    def export_to_pdf(self):
        """Export basic data to PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.pdf'
            filepath = os.path.join(self.output_dir, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title = Paragraph("Cloud Migration Plan", styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))
            
            # Content
            content = [
                "Server Inventory:",
                "• web-server-01: 4 CPU, 16GB RAM, 100GB Storage",
                "• db-server-01: 8 CPU, 32GB RAM, 500GB Storage", 
                "• app-server-01: 6 CPU, 24GB RAM, 200GB Storage",
                "",
                "Summary:",
                "• Total Servers: 3",
                "• Total RAM: 72GB",
                "• Total Storage: 800GB",
                "• Estimated Monthly Cost: $2,450",
                "• Migration Timeline: 3 months"
            ]
            
            for line in content:
                p = Paragraph(line, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 6))
            
            doc.build(story)
            return filename
            
        except Exception as e:
            raise Exception(f"PDF export failed: {str(e)}")
    
    def export_to_word(self):
        """Export basic data to Word document"""
        try:
            from docx import Document
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.docx'
            filepath = os.path.join(self.output_dir, filename)
            
            doc = Document()
            doc.add_heading('Cloud Migration Plan', 0)
            
            # Server inventory section
            doc.add_heading('Server Inventory', level=1)
            servers = [
                'web-server-01: 4 CPU, 16GB RAM, 100GB Storage',
                'db-server-01: 8 CPU, 32GB RAM, 500GB Storage',
                'app-server-01: 6 CPU, 24GB RAM, 200GB Storage'
            ]
            
            for server in servers:
                p = doc.add_paragraph(server, style='List Bullet')
            
            # Summary section
            doc.add_heading('Summary', level=1)
            summary_items = [
                'Total Servers: 3',
                'Total RAM: 72GB', 
                'Total Storage: 800GB',
                'Estimated Monthly Cost: $2,450',
                'Migration Timeline: 3 months'
            ]
            
            for item in summary_items:
                doc.add_paragraph(item, style='List Bullet')
            
            doc.save(filepath)
            return filename
            
        except Exception as e:
            raise Exception(f"Word export failed: {str(e)}")
