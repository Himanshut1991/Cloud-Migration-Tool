import os
import pandas as pd
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

class ExportService:
    """Export migration plans to various formats"""
    
    def __init__(self, db, models=None):
        self.db = db
        self.models = models or {}
        self.output_dir = os.path.join(os.path.dirname(__file__), '..', 'exports')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_to_excel(self):
        """Export migration plan to Excel format"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.xlsx'
            filepath = os.path.join(self.output_dir, filename)
            
            # Create workbook
            wb = Workbook()
            
            # Remove default sheet
            wb.remove(wb.active)
            
            # Create sheets
            self._create_summary_sheet(wb)
            self._create_inventory_sheets(wb)
            self._create_cost_analysis_sheet(wb)
            self._create_timeline_sheet(wb)
            self._create_resource_plan_sheet(wb)
            
            # Save workbook
            wb.save(filepath)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to export to Excel: {str(e)}")
    
    def export_to_pdf(self):
        """Export migration plan to PDF format"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.pdf'
            filepath = os.path.join(self.output_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            story.append(Paragraph("Cloud Migration Plan", title_style))
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            summary_data = self._get_summary_data()
            story.append(Paragraph(f"Total Servers: {summary_data['servers_count']}", styles['Normal']))
            story.append(Paragraph(f"Total Databases: {summary_data['databases_count']}", styles['Normal']))
            story.append(Paragraph(f"Total File Shares: {summary_data['file_shares_count']}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Cost Analysis
            story.append(Paragraph("Cost Analysis", styles['Heading2']))
            cost_data = self._get_cost_summary()
            cost_table_data = [
                ['Cost Category', 'Amount (USD)'],
                ['Cloud Infrastructure (Annual)', f"${cost_data.get('annual_cloud_cost', 0):,.2f}"],
                ['Migration Services (One-time)', f"${cost_data.get('migration_services_cost', 0):,.2f}"],
                ['Total First Year Cost', f"${cost_data.get('total_first_year', 0):,.2f}"]
            ]
            cost_table = Table(cost_table_data)
            cost_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(cost_table)
            story.append(Spacer(1, 20))
            
            # Timeline
            story.append(Paragraph("Migration Timeline", styles['Heading2']))
            timeline_data = self._get_timeline_summary()
            story.append(Paragraph(f"Total Duration: {timeline_data.get('total_weeks', 0)} weeks", styles['Normal']))
            story.append(Paragraph(f"Key Milestones: {len(timeline_data.get('milestones', []))} major milestones", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Inventory Summary Tables
            # self._add_inventory_tables_to_pdf(story, styles)  # Commented out to avoid model access issues
            
            # Build PDF
            doc.build(story)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to export to PDF: {str(e)}")
    
    def export_to_word(self):
        """Export migration plan to Word document"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.docx'
            filepath = os.path.join(self.output_dir, filename)
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading('Cloud Migration Plan', 0)
            title.alignment = 1  # Center alignment
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_data = self._get_summary_data()
            
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"This migration plan covers the migration of {summary_data['servers_count']} servers, ")
            summary_para.add_run(f"{summary_data['databases_count']} databases, and {summary_data['file_shares_count']} file shares ")
            summary_para.add_run(f"to the cloud. The total data volume to be migrated is {summary_data.get('total_data_gb', 0):,.0f} GB.")
            
            # Cost Analysis
            doc.add_heading('Cost Analysis', level=1)
            cost_data = self._get_cost_summary()
            
            # Cost table
            cost_table = doc.add_table(rows=4, cols=2)
            cost_table.style = 'Table Grid'
            
            # Header row
            hdr_cells = cost_table.rows[0].cells
            hdr_cells[0].text = 'Cost Category'
            hdr_cells[1].text = 'Amount (USD)'
            
            # Data rows
            cost_table.cell(1, 0).text = 'Cloud Infrastructure (Annual)'
            cost_table.cell(1, 1).text = f"${cost_data.get('annual_cloud_cost', 0):,.2f}"
            cost_table.cell(2, 0).text = 'Migration Services (One-time)'
            cost_table.cell(2, 1).text = f"${cost_data.get('migration_services_cost', 0):,.2f}"
            cost_table.cell(3, 0).text = 'Total First Year Cost'
            cost_table.cell(3, 1).text = f"${cost_data.get('total_first_year', 0):,.2f}"
            
            # Timeline
            doc.add_heading('Migration Timeline', level=1)
            timeline_data = self._get_timeline_summary()
            
            timeline_para = doc.add_paragraph()
            timeline_para.add_run(f"The migration is planned to complete in {timeline_data.get('total_weeks', 0)} weeks, ")
            timeline_para.add_run(f"with {len(timeline_data.get('milestones', []))} major milestones.")
            
            # Inventory Details
            self._add_inventory_sections_to_word(doc)
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            recommendations_para = doc.add_paragraph()
            recommendations_para.add_run("Based on the analysis, we recommend a phased approach to migration ")
            recommendations_para.add_run("starting with non-critical systems and progressing to mission-critical workloads. ")
            recommendations_para.add_run("This approach minimizes risk and allows for lessons learned to be applied to subsequent phases.")
            
            # Save document
            doc.save(filepath)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to export to Word: {str(e)}")
    
    def _create_summary_sheet(self, wb):
        """Create executive summary sheet"""
        ws = wb.create_sheet("Executive Summary")
        
        # Headers
        headers = ['Metric', 'Value']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Summary data
        summary_data = self._get_summary_data()
        metrics = [
            ('Total Servers', summary_data['servers_count']),
            ('Total Databases', summary_data['databases_count']),
            ('Total File Shares', summary_data['file_shares_count']),
            ('Total Data Size (GB)', f"{summary_data.get('total_data_gb', 0):,.0f}"),
            ('Report Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for row, (metric, value) in enumerate(metrics, 2):
            ws.cell(row=row, column=1, value=metric)
            ws.cell(row=row, column=2, value=value)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def _create_inventory_sheets(self, wb):
        """Create inventory sheets for servers, databases, and file shares"""
        
        # Servers sheet
        servers_ws = wb.create_sheet("Server Inventory")
        Server = self.models.get('Server')
        
        if Server:
            servers = Server.query.all()
            if servers:
                server_headers = ['Server ID', 'OS Type', 'vCPU', 'RAM (GB)', 'Disk Size (GB)', 'Disk Type', 'Technologies']
                for col, header in enumerate(server_headers, 1):
                    cell = servers_ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                
                for row, server in enumerate(servers, 2):
                    servers_ws.cell(row=row, column=1, value=server.server_id)
                    servers_ws.cell(row=row, column=2, value=server.os_type)
                    servers_ws.cell(row=row, column=3, value=server.vcpu)
                    servers_ws.cell(row=row, column=4, value=server.ram)
                    servers_ws.cell(row=row, column=5, value=server.disk_size)
                    servers_ws.cell(row=row, column=6, value=server.disk_type)
                    servers_ws.cell(row=row, column=7, value=server.technology or '')
        
        # Databases sheet
        db_ws = wb.create_sheet("Database Inventory")
        Database = self.models.get('Database')
        
        if Database:
            databases = Database.query.all()
            if databases:
                db_headers = ['DB Name', 'DB Type', 'Size (GB)', 'HA/DR', 'Backup Frequency', 'Licensing']
                for col, header in enumerate(db_headers, 1):
                    cell = db_ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                
                for row, database in enumerate(databases, 2):
                    db_ws.cell(row=row, column=1, value=database.db_name)
                    db_ws.cell(row=row, column=2, value=database.db_type)
                    db_ws.cell(row=row, column=3, value=database.size_gb)
                    db_ws.cell(row=row, column=4, value='Yes' if database.ha_dr_required else 'No')
                    db_ws.cell(row=row, column=5, value=database.backup_frequency)
                    db_ws.cell(row=row, column=6, value=database.licensing_model)
        
        # File Shares sheet
        fs_ws = wb.create_sheet("File Share Inventory")
        FileShare = self.models.get('FileShare')
        
        if FileShare:
            file_shares = FileShare.query.all()
            if file_shares:
                fs_headers = ['Share Name', 'Size (GB)', 'Access Pattern', 'Snapshots', 'Retention (Days)']
                for col, header in enumerate(fs_headers, 1):
                    cell = fs_ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
                
                for row, file_share in enumerate(file_shares, 2):
                    fs_ws.cell(row=row, column=1, value=file_share.share_name)
                    fs_ws.cell(row=row, column=2, value=file_share.total_size_gb)
                    fs_ws.cell(row=row, column=3, value=file_share.access_pattern)
                    fs_ws.cell(row=row, column=4, value='Yes' if file_share.snapshot_required else 'No')
                    fs_ws.cell(row=row, column=5, value=file_share.retention_days)
                fs_ws.cell(row=row, column=1, value=file_share.share_name)
                fs_ws.cell(row=row, column=2, value=file_share.total_size_gb)
                fs_ws.cell(row=row, column=3, value=file_share.access_pattern)
                fs_ws.cell(row=row, column=4, value='Yes' if file_share.snapshot_required else 'No')
                fs_ws.cell(row=row, column=5, value=file_share.retention_days)
    
    def _create_cost_analysis_sheet(self, wb):
        """Create cost analysis sheet"""
        ws = wb.create_sheet("Cost Analysis")
        
        # This would integrate with the CostCalculator service
        # For now, creating a basic structure
        headers = ['Cost Category', 'Monthly Cost', 'Annual Cost']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Sample cost data - would be replaced with actual calculations
        cost_data = [
            ('Server Migration', 1000, 12000),
            ('Database Migration', 500, 6000),
            ('Storage Migration', 200, 2400),
            ('Professional Services', 15000, 15000)  # One-time cost
        ]
        
        for row, (category, monthly, annual) in enumerate(cost_data, 2):
            ws.cell(row=row, column=1, value=category)
            ws.cell(row=row, column=2, value=monthly)
            ws.cell(row=row, column=3, value=annual)
    
    def _create_timeline_sheet(self, wb):
        """Create timeline sheet"""
        ws = wb.create_sheet("Timeline")
        
        headers = ['Phase', 'Start Date', 'End Date', 'Duration (Weeks)', 'Status']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Sample timeline data
        timeline_data = [
            ('Assessment and Planning', '2025-01-01', '2025-01-14', 2, 'Not Started'),
            ('Environment Setup', '2025-01-15', '2025-01-28', 2, 'Not Started'),
            ('Pilot Migration', '2025-01-29', '2025-02-18', 3, 'Not Started'),
            ('Data Migration', '2025-02-19', '2025-03-18', 4, 'Not Started'),
            ('Server Migration', '2025-03-19', '2025-04-15', 4, 'Not Started'),
            ('Testing and Validation', '2025-04-16', '2025-05-06', 3, 'Not Started'),
            ('Cutover and Go-Live', '2025-05-07', '2025-05-20', 2, 'Not Started'),
            ('Post-Migration Support', '2025-05-21', '2025-06-17', 4, 'Not Started')
        ]
        
        for row, (phase, start, end, duration, status) in enumerate(timeline_data, 2):
            ws.cell(row=row, column=1, value=phase)
            ws.cell(row=row, column=2, value=start)
            ws.cell(row=row, column=3, value=end)
            ws.cell(row=row, column=4, value=duration)
            ws.cell(row=row, column=5, value=status)
    
    def _create_resource_plan_sheet(self, wb):
        """Create resource planning sheet"""
        ws = wb.create_sheet("Resource Plan")
        
        headers = ['Role', 'Duration (Weeks)', 'Hours/Week', 'Rate/Hour', 'Total Cost']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        resource_rates = ResourceRate.query.all()
        for row, rate in enumerate(resource_rates, 2):
            total_cost = rate.duration_weeks * rate.hours_per_week * rate.rate_per_hour
            ws.cell(row=row, column=1, value=rate.role)
            ws.cell(row=row, column=2, value=rate.duration_weeks)
            ws.cell(row=row, column=3, value=rate.hours_per_week)
            ws.cell(row=row, column=4, value=rate.rate_per_hour)
            ws.cell(row=row, column=5, value=total_cost)
    
    def _get_summary_data(self):
        """Get summary data for reports"""
        Server = self.models.get('Server')
        Database = self.models.get('Database')
        FileShare = self.models.get('FileShare')
        
        servers_count = Server.query.count() if Server else 0
        databases_count = Database.query.count() if Database else 0
        file_shares_count = FileShare.query.count() if FileShare else 0
        
        total_db_size = 0
        total_fs_size = 0
        
        if Database:
            total_db_size = self.db.session.query(self.db.func.sum(Database.size_gb)).scalar() or 0
        if FileShare:
            total_fs_size = self.db.session.query(self.db.func.sum(FileShare.total_size_gb)).scalar() or 0
        
        return {
            'servers_count': servers_count,
            'databases_count': databases_count,
            'file_shares_count': file_shares_count,
            'total_data_gb': total_db_size + total_fs_size
        }
    
    def _get_cost_summary(self):
        """Get cost summary data"""
        # This would integrate with actual cost calculations
        return {
            'annual_cloud_cost': 50000,
            'migration_services_cost': 75000,
            'total_first_year': 125000
        }
    
    def _get_timeline_summary(self):
        """Get timeline summary data"""
        return {
            'total_weeks': 24,
            'milestones': [
                'Assessment Complete',
                'Environment Ready',
                'Pilot Complete',
                'Data Migration Complete',
                'Go-Live'
            ]
        }
    
    def _add_inventory_tables_to_pdf(self, story, styles):
        """Add inventory tables to PDF"""
        # Server inventory
        story.append(Paragraph("Server Inventory", styles['Heading2']))
        Server = self.models.get('Server')
        if Server:
            servers = Server.query.all()
            if servers:
                server_data = [['Server ID', 'OS', 'vCPU', 'RAM (GB)', 'Disk (GB)']]
                for server in servers[:10]:  # Limit to first 10 for space
                    server_data.append([
                        server.server_id,
                        server.os_type,
                        str(server.vcpu),
                        str(server.ram),
                        str(server.disk_size)
                    ])
                
                server_table = Table(server_data)
                server_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(server_table)
        
        story.append(Spacer(1, 20))
    
    def _add_inventory_sections_to_word(self, doc):
        """Add inventory sections to Word document"""
        # Server inventory
        doc.add_heading('Server Inventory', level=2)
        Server = self.models.get('Server')
        if Server:
            servers = Server.query.all()
            if servers:
                server_table = doc.add_table(rows=1, cols=5)
                server_table.style = 'Table Grid'
                
                # Header
                hdr_cells = server_table.rows[0].cells
                hdr_cells[0].text = 'Server ID'
                hdr_cells[1].text = 'OS Type'
                hdr_cells[2].text = 'vCPU'
                hdr_cells[3].text = 'RAM (GB)'
                hdr_cells[4].text = 'Disk (GB)'
                
                # Data rows (limit to first 10)
                for server in servers[:10]:
                    row_cells = server_table.add_row().cells
                    row_cells[0].text = server.server_id
                    row_cells[1].text = server.os_type
                    row_cells[2].text = str(server.vcpu)
                    row_cells[3].text = str(server.ram)
                    row_cells[4].text = str(server.disk_size)
