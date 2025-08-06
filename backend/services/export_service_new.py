import os
import pandas as pd
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

class ExportService:
    """Export migration plans to various formats"""
    
    def __init__(self, db, models=None):
        self.db = db
        self.models = models or {}
        self.output_dir = os.path.join(os.path.dirname(__file__), '..', 'exports')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_to_excel(self):
        """Export migration plan to Excel format"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.xlsx'
            filepath = os.path.join(self.output_dir, filename)
            
            # Create workbook with multiple sheets
            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # Create summary sheet
            self._create_summary_sheet(wb)
            
            # Create inventory sheets
            self._create_servers_sheet(wb)
            self._create_databases_sheet(wb)
            self._create_file_shares_sheet(wb)
            
            # Create analysis sheets
            self._create_cost_analysis_sheet(wb)
            self._create_timeline_sheet(wb)
            
            wb.save(filepath)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to export to Excel: {str(e)}")
    
    def export_to_pdf(self):
        """Export comprehensive migration plan to PDF format"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.pdf'
            filepath = os.path.join(self.output_dir, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1,
                textColor=colors.darkblue
            )
            
            # Title Page
            story.append(Paragraph("CLOUD MIGRATION PLAN", title_style))
            story.append(Spacer(1, 20))
            story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
            story.append(Spacer(1, 40))
            
            # Executive Summary
            story.append(Paragraph("EXECUTIVE SUMMARY", styles['Heading2']))
            summary_data = self._get_summary_data()
            
            story.append(Paragraph(f"This comprehensive migration plan covers the migration of {summary_data['servers_count']} servers, {summary_data['databases_count']} databases, and {summary_data['file_shares_count']} file shares to the cloud platform.", styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Infrastructure Overview Table
            infra_data = [
                ['Infrastructure Type', 'Count', 'Total Size (GB)', 'Migration Priority'],
                ['Servers', str(summary_data['servers_count']), 'N/A', 'High'],
                ['Databases', str(summary_data['databases_count']), str(summary_data.get('total_db_size', 0)), 'Critical'],
                ['File Shares', str(summary_data['file_shares_count']), str(summary_data.get('total_fs_size', 0)), 'Medium'],
                ['Total Data Volume', '-', f"{summary_data['total_data_gb']:,.0f}", '-']
            ]
            
            infra_table = Table(infra_data)
            infra_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))
            story.append(infra_table)
            story.append(Spacer(1, 20))
            
            # Detailed Server Inventory
            story.append(Paragraph("SERVER INVENTORY", styles['Heading2']))
            Server = self.models.get('Server')
            if Server:
                servers = Server.query.limit(20).all()
                if servers:
                    server_data = [['Server ID', 'OS Type', 'vCPU', 'RAM (GB)', 'Disk (GB)', 'Hosting', 'Technology']]
                    for server in servers:
                        server_data.append([
                            f"Server-{server.server_id}",
                            server.os_type or 'N/A',
                            str(server.vcpu or 0),
                            str(server.ram or 0),
                            str(server.disk_size or 0),
                            server.current_hosting or 'On-Premises',
                            server.technology or 'General'
                        ])
                    
                    server_table = Table(server_data)
                    server_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
                    ]))
                    story.append(server_table)
            story.append(Spacer(1, 20))
            
            # Database Inventory
            story.append(Paragraph("DATABASE INVENTORY", styles['Heading2']))
            Database = self.models.get('Database')
            if Database:
                databases = Database.query.limit(15).all()
                if databases:
                    db_data = [['Database Name', 'Type', 'Size (GB)', 'HA/DR', 'Performance Tier', 'Backup Freq']]
                    for db in databases:
                        db_data.append([
                            db.db_name or f"DB-{db.id}",
                            db.db_type or 'SQL Server',
                            str(db.size_gb or 0),
                            'Yes' if db.ha_dr_required else 'No',
                            db.performance_tier or 'Standard',
                            db.backup_frequency or 'Daily'
                        ])
                    
                    db_table = Table(db_data)
                    db_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
                    ]))
                    story.append(db_table)
            story.append(Spacer(1, 20))
            
            # Cost Analysis
            story.append(Paragraph("COST ANALYSIS", styles['Heading2']))
            cost_data = self._get_cost_summary()
            
            cost_table_data = [
                ['Cost Category', 'Monthly Cost', 'Annual Cost', 'Notes'],
                ['EC2 Instances', f"${cost_data.get('monthly_compute', 2500):,.2f}", f"${cost_data.get('annual_compute', 30000):,.2f}", 'Reserved instances recommended'],
                ['RDS Databases', f"${cost_data.get('monthly_database', 1200):,.2f}", f"${cost_data.get('annual_database', 14400):,.2f}", 'Multi-AZ for HA'],
                ['Storage (S3/EBS)', f"${cost_data.get('monthly_storage', 800):,.2f}", f"${cost_data.get('annual_storage', 9600):,.2f}", 'Intelligent tiering'],
                ['Data Transfer', f"${cost_data.get('monthly_transfer', 300):,.2f}", f"${cost_data.get('annual_transfer', 3600):,.2f}", 'CloudFront recommended'],
                ['Migration Services', f"${cost_data.get('monthly_migration', 0):,.2f}", f"${cost_data.get('migration_services_cost', 25000):,.2f}", 'One-time cost'],
                ['TOTAL FIRST YEAR', f"${cost_data.get('total_monthly', 4800):,.2f}", f"${cost_data.get('total_first_year', 82600):,.2f}", 'Including migration']
            ]
            
            cost_table = Table(cost_table_data)
            cost_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, -1), (-1, -1), colors.lightgreen),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.lightgrey])
            ]))
            story.append(cost_table)
            story.append(Spacer(1, 20))
            
            # Migration Timeline
            story.append(Paragraph("MIGRATION TIMELINE", styles['Heading2']))
            timeline_data = self._get_timeline_summary()
            
            timeline_table_data = [
                ['Phase', 'Duration', 'Start Date', 'Activities', 'Dependencies'],
                ['Assessment', '4 weeks', '2024-03-01', 'Inventory, Analysis, Planning', 'None'],
                ['Infrastructure Setup', '3 weeks', '2024-03-29', 'AWS Account, Networking, Security', 'Assessment Complete'],
                ['Data Migration', f"{timeline_data.get('data_weeks', 6)} weeks", '2024-04-19', 'Database Migration, File Transfer', 'Infrastructure Ready'],
                ['Application Migration', f"{timeline_data.get('app_weeks', 8)} weeks", '2024-05-31', 'Server Migration, Testing', 'Data Migration 80%'],
                ['Go-Live & Optimization', '2 weeks', '2024-07-26', 'Cutover, Optimization, Support', 'All Migration Complete'],
            ]
            
            timeline_table = Table(timeline_table_data)
            timeline_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.darkred),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(timeline_table)
            story.append(Spacer(1, 20))
            
            # Risk Assessment
            story.append(Paragraph("RISK ASSESSMENT & MITIGATION", styles['Heading2']))
            
            risks_data = [
                ['Risk Category', 'Probability', 'Impact', 'Mitigation Strategy'],
                ['Data Loss During Migration', 'Low', 'High', 'Full backup, incremental sync, rollback plan'],
                ['Application Downtime', 'Medium', 'High', 'Blue-green deployment, maintenance windows'],
                ['Performance Degradation', 'Medium', 'Medium', 'Load testing, auto-scaling, monitoring'],
                ['Security Vulnerabilities', 'Low', 'High', 'Security assessment, IAM policies, encryption'],
                ['Budget Overrun', 'Medium', 'Medium', 'Reserved instances, cost monitoring, alerts'],
                ['Timeline Delays', 'High', 'Medium', 'Buffer time, parallel workstreams, risk tracking']
            ]
            
            risks_table = Table(risks_data)
            risks_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.orange),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            story.append(risks_table)
            
            doc.build(story)
            return filepath
            
        except Exception as e:
            raise Exception(f"Failed to export to PDF: {str(e)}")
    
    def export_to_word(self):
        """Export comprehensive migration plan to Word document"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.docx'
            filepath = os.path.join(self.output_dir, filename)
            
            print(f"📄 Creating comprehensive Word document: {filename}")
            
            doc = Document()
            
            # Title Page
            title = doc.add_heading('Cloud Migration Assessment Report', 0)
            title.alignment = 1
            
            # Add generation date
            date_para = doc.add_paragraph()
            date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            date_para.alignment = 1
            
            doc.add_page_break()
            
            # Executive Summary
            print("📝 Adding executive summary...")
            try:
                self._add_executive_summary_word(doc)
                print("✅ Executive summary added")
            except Exception as e:
                print(f"❌ Error in executive summary: {e}")
                raise
            
            # Infrastructure Overview
            print("📋 Adding infrastructure overview...")
            try:
                self._add_infrastructure_overview_word(doc)
                print("✅ Infrastructure overview added")
            except Exception as e:
                print(f"❌ Error in infrastructure overview: {e}")
                raise
            
            # Detailed Inventory
            print("📊 Adding detailed inventory...")
            try:
                self._add_detailed_inventory_word(doc)
                print("✅ Detailed inventory added")
            except Exception as e:
                print(f"❌ Error in detailed inventory: {e}")
                raise
            
            # Cost Analysis
            print("💰 Adding cost analysis...")
            try:
                self._add_cost_analysis_word(doc)
                print("✅ Cost analysis added")
            except Exception as e:
                print(f"❌ Error in cost analysis: {e}")
                raise
            
            # Migration Timeline
            print("📅 Adding migration timeline...")
            try:
                self._add_timeline_word(doc)
                print("✅ Migration timeline added")
            except Exception as e:
                print(f"❌ Error in migration timeline: {e}")
                raise
            
            # Risk Assessment
            print("⚠️ Adding risk assessment...")
            try:
                self._add_risk_assessment_word(doc)
                print("✅ Risk assessment added")
            except Exception as e:
                print(f"❌ Error in risk assessment: {e}")
                raise
            
            # Recommendations
            print("💡 Adding recommendations...")
            try:
                self._add_recommendations_word(doc)
                print("✅ Recommendations added")
            except Exception as e:
                print(f"❌ Error in recommendations: {e}")
                raise
            
            print("💾 Saving document...")
            doc.save(filepath)
            print(f"✅ Comprehensive Word document saved: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"❌ Word export error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Failed to export to Word: {str(e)}")
    
    def _add_executive_summary_word(self, doc):
        """Add comprehensive executive summary to Word document"""
        doc.add_heading('Executive Summary', level=1)
        
        # Get data
        print("🔍 Getting summary data...")
        summary_data = self._get_summary_data()
        print(f"📊 Summary data: {summary_data}")
        
        # Overview paragraph
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"This comprehensive migration assessment analyzes the migration of ")
        summary_para.add_run(f"{summary_data['servers_count']} servers, {summary_data['databases_count']} databases, ")
        summary_para.add_run(f"and {summary_data['file_shares_count']} file shares to the cloud. ")
        summary_para.add_run(f"The total infrastructure includes {summary_data.get('total_vcpu', 0)} vCPUs, ")
        summary_para.add_run(f"{summary_data.get('total_ram_gb', 0):,.0f} GB RAM, and ")
        summary_para.add_run(f"{summary_data.get('total_data_gb', 0):,.0f} GB of data storage.")
        
        # Key metrics table
        print("📋 Creating metrics table...")
        metrics_table = doc.add_table(rows=7, cols=2)
        metrics_table.style = 'Table Grid'
        
        # Headers
        hdr_cells = metrics_table.rows[0].cells
        hdr_cells[0].text = 'Migration Metric'
        hdr_cells[1].text = 'Value'
        
        # Key metrics data
        print("📊 Preparing metrics data...")
        metrics_data = [
            ('Total Infrastructure Components', str(summary_data['servers_count'] + summary_data['databases_count'] + summary_data['file_shares_count'])),
            ('Estimated Migration Duration', f"{summary_data.get('estimated_weeks', 12)} weeks"),
            ('Migration Complexity', summary_data.get('complexity_level', 'Medium')),
            ('Estimated Monthly Cloud Cost', f"${summary_data.get('monthly_cost', 5000):,.2f}"),
            ('Estimated Annual Cloud Cost', f"${summary_data.get('annual_cost', 60000):,.2f}"),
            ('Primary Migration Strategy', summary_data.get('primary_strategy', 'Rehost (Lift & Shift)'))
        ]
        
        print(f"📋 Metrics data prepared: {len(metrics_data)} items")
        print(f"📊 Table has {len(metrics_table.rows)} rows")
        
        # Populate table
        print("📝 Populating table...")
        try:
            for i, (metric, value) in enumerate(metrics_data, 1):
                print(f"Adding row {i}: {metric} = {value}")
                if i < len(metrics_table.rows):
                    metrics_table.rows[i].cells[0].text = str(metric)
                    metrics_table.rows[i].cells[1].text = str(value)
                else:
                    print(f"⚠️ Row {i} exceeds table size")
            print("✅ Table populated successfully")
        except Exception as e:
            print(f"❌ Error populating table: {e}")
            raise
    
    def _add_infrastructure_overview_word(self, doc):
        """Add infrastructure overview to Word document"""
        doc.add_heading('Infrastructure Overview', level=1)
        
        # Current state analysis
        doc.add_heading('Current State Analysis', level=2)
        current_para = doc.add_paragraph()
        current_para.add_run("The existing infrastructure consists of a mixed environment with ")
        current_para.add_run("various operating systems, database platforms, and storage solutions. ")
        current_para.add_run("This diversity presents both opportunities for optimization and challenges for migration planning.")
        
        # Resource breakdown
        doc.add_heading('Resource Breakdown', level=2)
        
        # Get servers for analysis
        Server = self.models.get('Server')
        if Server:
            servers = Server.query.all()
            
            # OS distribution
            os_counts = {}
            total_vcpu = 0
            total_ram = 0
            total_disk = 0
            
            for server in servers:
                os_type = server.os_type or 'Unknown'
                os_counts[os_type] = os_counts.get(os_type, 0) + 1
                total_vcpu += server.vcpu or 0
                total_ram += server.ram or 0
                total_disk += server.disk_size or 0
            
            # Create OS distribution table
            if os_counts:
                doc.add_heading('Operating System Distribution', level=3)
                os_table = doc.add_table(rows=len(os_counts) + 1, cols=3)
                os_table.style = 'Table Grid'
                
                # Headers
                os_table.rows[0].cells[0].text = 'Operating System'
                os_table.rows[0].cells[1].text = 'Count'
                os_table.rows[0].cells[2].text = 'Percentage'
                
                total_servers = sum(os_counts.values())
                for i, (os_type, count) in enumerate(os_counts.items(), 1):
                    os_table.rows[i].cells[0].text = os_type
                    os_table.rows[i].cells[1].text = str(count)
                    os_table.rows[i].cells[2].text = f"{(count/total_servers)*100:.1f}%"
        
        # Database analysis
        Database = self.models.get('Database')
        if Database:
            databases = Database.query.all()
            
            doc.add_heading('Database Environment', level=2)
            db_para = doc.add_paragraph()
            db_para.add_run(f"The database environment includes {len(databases)} database instances ")
            db_para.add_run("across various database management systems. Each database has been ")
            db_para.add_run("analyzed for migration complexity, compatibility, and optimization opportunities.")
    
    def _add_detailed_inventory_word(self, doc):
        """Add detailed inventory tables to Word document"""
        doc.add_heading('Detailed Infrastructure Inventory', level=1)
        
        # Server inventory
        doc.add_heading('Server Inventory', level=2)
        Server = self.models.get('Server')
        if Server:
            servers = Server.query.limit(20).all()  # Limit for document readability
            
            if servers:
                server_table = doc.add_table(rows=len(servers) + 1, cols=6)
                server_table.style = 'Table Grid'
                
                # Headers
                headers = ['Server Name', 'OS Type', 'vCPU', 'RAM (GB)', 'Disk (GB)', 'Environment']
                for i, header in enumerate(headers):
                    server_table.rows[0].cells[i].text = header
                
                # Server data
                for i, server in enumerate(servers, 1):
                    server_table.rows[i].cells[0].text = getattr(server, 'id', f'Server-{i}')
                    server_table.rows[i].cells[1].text = getattr(server, 'os_type', 'Unknown')
                    server_table.rows[i].cells[2].text = str(getattr(server, 'vcpu', 0))
                    server_table.rows[i].cells[3].text = str(getattr(server, 'ram', 0))
                    server_table.rows[i].cells[4].text = str(getattr(server, 'disk_size', 0))
                    server_table.rows[i].cells[5].text = getattr(server, 'environment', 'Production')
                
                if len(Server.query.all()) > 20:
                    doc.add_paragraph(f"Note: Showing first 20 servers. Total servers: {Server.query.count()}")
        
        # Database inventory
        doc.add_heading('Database Inventory', level=2)
        Database = self.models.get('Database')
        if Database:
            databases = Database.query.limit(15).all()
            
            if databases:
                db_table = doc.add_table(rows=len(databases) + 1, cols=5)
                db_table.style = 'Table Grid'
                
                # Headers
                headers = ['Database Name', 'Type', 'Version', 'Size (GB)', 'Environment']
                for i, header in enumerate(headers):
                    db_table.rows[0].cells[i].text = header
                
                # Database data
                for i, db in enumerate(databases, 1):
                    db_table.rows[i].cells[0].text = getattr(db, 'database_name', f'DB-{i}')
                    db_table.rows[i].cells[1].text = getattr(db, 'db_type', 'Unknown')
                    db_table.rows[i].cells[2].text = getattr(db, 'version', 'N/A')
                    db_table.rows[i].cells[3].text = str(getattr(db, 'size_gb', 0))
                    db_table.rows[i].cells[4].text = getattr(db, 'environment', 'Production')
                
                if len(Database.query.all()) > 15:
                    doc.add_paragraph(f"Note: Showing first 15 databases. Total databases: {Database.query.count()}")
    
    def _add_cost_analysis_word(self, doc):
        """Add comprehensive cost analysis to Word document"""
        doc.add_heading('Cost Analysis', level=1)
        
        cost_data = self._get_cost_summary()
        
        # Overview paragraph
        cost_para = doc.add_paragraph()
        cost_para.add_run("The cost analysis provides estimated expenses for cloud migration and ongoing operations. ")
        cost_para.add_run("These estimates are based on current infrastructure sizing and standard cloud pricing models.")
        
        # Cost breakdown table
        doc.add_heading('Cost Breakdown', level=2)
        cost_table = doc.add_table(rows=6, cols=3)
        cost_table.style = 'Table Grid'
        
        # Headers
        cost_table.rows[0].cells[0].text = 'Cost Category'
        cost_table.rows[0].cells[1].text = 'Monthly Cost'
        cost_table.rows[0].cells[2].text = 'Annual Cost'
        
        # Cost data
        cost_items = [
            ('Compute Services', f"${cost_data.get('monthly_compute', 2000):,.2f}", f"${cost_data.get('annual_compute', 24000):,.2f}"),
            ('Storage Services', f"${cost_data.get('monthly_storage', 500):,.2f}", f"${cost_data.get('annual_storage', 6000):,.2f}"),
            ('Database Services', f"${cost_data.get('monthly_database', 1500):,.2f}", f"${cost_data.get('annual_database', 18000):,.2f}"),
            ('Network & Security', f"${cost_data.get('monthly_network', 300):,.2f}", f"${cost_data.get('annual_network', 3600):,.2f}"),
            ('Total Operational Cost', f"${cost_data.get('total_monthly', 4300):,.2f}", f"${cost_data.get('total_annual', 51600):,.2f}")
        ]
        
        for i, (category, monthly, annual) in enumerate(cost_items, 1):
            cost_table.rows[i].cells[0].text = category
            cost_table.rows[i].cells[1].text = monthly
            cost_table.rows[i].cells[2].text = annual
        
        # Migration costs
        doc.add_heading('One-Time Migration Costs', level=2)
        migration_para = doc.add_paragraph()
        migration_para.add_run(f"Professional Services: ${cost_data.get('professional_services', 25000):,.2f}")
        migration_para.add_run(f"\\nData Transfer: ${cost_data.get('data_transfer', 5000):,.2f}")
        migration_para.add_run(f"\\nTesting & Validation: ${cost_data.get('testing_validation', 10000):,.2f}")
        migration_para.add_run(f"\\nTotal Migration Cost: ${cost_data.get('total_migration', 40000):,.2f}")
    
    def _add_timeline_word(self, doc):
        """Add migration timeline to Word document"""
        doc.add_heading('Migration Timeline', level=1)
        
        timeline_data = self._get_timeline_summary()
        
        # Timeline overview
        timeline_para = doc.add_paragraph()
        timeline_para.add_run(f"The migration is planned to span {timeline_data.get('total_weeks', 16)} weeks, ")
        timeline_para.add_run("organized into distinct phases to minimize risk and ensure business continuity.")
        
        # Phase breakdown
        doc.add_heading('Migration Phases', level=2)
        phases = [
            ('Phase 1: Assessment & Planning', '2 weeks', 'Infrastructure discovery, application mapping, migration planning'),
            ('Phase 2: Environment Setup', '3 weeks', 'Cloud environment provisioning, network configuration, security setup'),
            ('Phase 3: Data Migration', '4 weeks', 'Database migration, file system transfer, data validation'),
            ('Phase 4: Application Migration', '5 weeks', 'Server migration, application deployment, configuration'),
            ('Phase 5: Testing & Validation', '2 weeks', 'End-to-end testing, performance validation, user acceptance'),
            ('Phase 6: Go-Live & Support', '1 week', 'Production cutover, monitoring setup, knowledge transfer')
        ]
        
        phase_table = doc.add_table(rows=len(phases) + 1, cols=3)
        phase_table.style = 'Table Grid'
        
        # Headers
        phase_table.rows[0].cells[0].text = 'Phase'
        phase_table.rows[0].cells[1].text = 'Duration'
        phase_table.rows[0].cells[2].text = 'Key Activities'
        
        for i, (phase, duration, activities) in enumerate(phases, 1):
            phase_table.rows[i].cells[0].text = phase
            phase_table.rows[i].cells[1].text = duration
            phase_table.rows[i].cells[2].text = activities
    
    def _add_risk_assessment_word(self, doc):
        """Add risk assessment to Word document"""
        doc.add_heading('Risk Assessment', level=1)
        
        # Risk overview
        risk_para = doc.add_paragraph()
        risk_para.add_run("The following risks have been identified and should be addressed during the migration planning process:")
        
        # Risk table
        risks = [
            ('Data Loss', 'High', 'Implement comprehensive backup and validation procedures'),
            ('Downtime Risk', 'Medium', 'Plan migration during maintenance windows, use blue-green deployment'),
            ('Performance Issues', 'Medium', 'Conduct thorough performance testing and monitoring'),
            ('Security Vulnerabilities', 'High', 'Implement security best practices and compliance measures'),
            ('Cost Overruns', 'Medium', 'Regular cost monitoring and optimization reviews'),
            ('Skill Gaps', 'Low', 'Provide training and documentation for operations teams')
        ]
        
        risk_table = doc.add_table(rows=len(risks) + 1, cols=3)
        risk_table.style = 'Table Grid'
        
        # Headers
        risk_table.rows[0].cells[0].text = 'Risk'
        risk_table.rows[0].cells[1].text = 'Severity'
        risk_table.rows[0].cells[2].text = 'Mitigation Strategy'
        
        for i, (risk, severity, mitigation) in enumerate(risks, 1):
            risk_table.rows[i].cells[0].text = risk
            risk_table.rows[i].cells[1].text = severity
            risk_table.rows[i].cells[2].text = mitigation
    
    def _add_recommendations_word(self, doc):
        """Add recommendations to Word document"""
        doc.add_heading('Recommendations', level=1)
        
        # Strategic recommendations
        doc.add_heading('Strategic Recommendations', level=2)
        strategic_para = doc.add_paragraph()
        strategic_para.add_run("Based on the infrastructure assessment, the following strategic recommendations are provided:")
        
        recommendations = [
            "Adopt a phased migration approach to minimize business impact",
            "Prioritize rehosting (lift-and-shift) for quick wins, then optimize",
            "Implement cloud-native monitoring and alerting solutions",
            "Establish governance processes for cloud resource management",
            "Consider reserved instances for predictable workloads to reduce costs",
            "Implement automated backup and disaster recovery solutions",
            "Establish a cloud center of excellence for ongoing optimization"
        ]
        
        for recommendation in recommendations:
            para = doc.add_paragraph(style='ListBullet')
            para.add_run(recommendation)
        
        # Technical recommendations
        doc.add_heading('Technical Recommendations', level=2)
        technical_para = doc.add_paragraph()
        technical_para.add_run("The following technical recommendations should be considered during implementation:")
        
        technical_recs = [
            "Use infrastructure as code (IaC) for consistent deployments",
            "Implement container orchestration where applicable",
            "Establish CI/CD pipelines for application deployments",
            "Use cloud-native databases where possible for better performance",
            "Implement auto-scaling for variable workloads",
            "Use content delivery networks (CDN) for improved performance"
        ]
        
        for rec in technical_recs:
            para = doc.add_paragraph(style='ListBullet')
            para.add_run(rec)
    
    # Helper methods
    def _create_summary_sheet(self, wb):
        """Create summary sheet in Excel"""
        ws = wb.create_sheet("Summary")
        
        # Header styling
        header_font = Font(bold=True, size=14)
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Title
        ws['A1'] = "Cloud Migration Plan Summary"
        ws['A1'].font = Font(bold=True, size=16)
        ws.merge_cells('A1:B1')
        
        # Summary data
        summary_data = self._get_summary_data()
        
        ws['A3'] = "Inventory Summary"
        ws['A3'].font = header_font
        ws['A3'].fill = header_fill
        
        ws['A4'] = "Servers"
        ws['B4'] = summary_data['servers_count']
        ws['A5'] = "Databases"
        ws['B5'] = summary_data['databases_count']
        ws['A6'] = "File Shares"
        ws['B6'] = summary_data['file_shares_count']
        ws['A7'] = "Total Data (GB)"
        ws['B7'] = summary_data['total_data_gb']
        
        # Cost summary
        cost_data = self._get_cost_summary()
        
        ws['A9'] = "Cost Summary"
        ws['A9'].font = header_font
        ws['A9'].fill = header_fill
        
        ws['A10'] = "Annual Cloud Cost"
        ws['B10'] = cost_data['annual_cloud_cost']
        ws['A11'] = "Migration Services"
        ws['B11'] = cost_data['migration_services_cost']
        ws['A12'] = "Total First Year"
        ws['B12'] = cost_data['total_first_year']
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def _create_servers_sheet(self, wb):
        """Create servers inventory sheet"""
        ws = wb.create_sheet("Servers")
        
        # Headers
        headers = ['Name', 'Environment', 'OS', 'CPU Cores', 'Memory (GB)', 'Storage (GB)', 'Status']
        for i, header in enumerate(headers, 1):
            ws.cell(row=1, column=i, value=header).font = Font(bold=True)
        
        # Get server data
        Server = self.models.get('Server')
        if Server:
            servers = Server.query.all()
            for i, server in enumerate(servers, 2):
                ws.cell(row=i, column=1, value=server.name)
                ws.cell(row=i, column=2, value=server.environment)
                ws.cell(row=i, column=3, value=server.operating_system)
                ws.cell(row=i, column=4, value=server.cpu_cores)
                ws.cell(row=i, column=5, value=server.memory_gb)
                ws.cell(row=i, column=6, value=server.storage_gb)
                ws.cell(row=i, column=7, value=server.status)
    
    def _create_databases_sheet(self, wb):
        """Create databases inventory sheet"""
        ws = wb.create_sheet("Databases")
        
        headers = ['Name', 'Type', 'Version', 'Size (GB)', 'Environment', 'Status']
        for i, header in enumerate(headers, 1):
            ws.cell(row=1, column=i, value=header).font = Font(bold=True)
        
        Database = self.models.get('Database')
        if Database:
            databases = Database.query.all()
            for i, db in enumerate(databases, 2):
                ws.cell(row=i, column=1, value=db.name)
                ws.cell(row=i, column=2, value=db.database_type)
                ws.cell(row=i, column=3, value=db.version)
                ws.cell(row=i, column=4, value=db.size_gb)
                ws.cell(row=i, column=5, value=db.environment)
                ws.cell(row=i, column=6, value=db.status)
    
    def _create_file_shares_sheet(self, wb):
        """Create file shares inventory sheet"""
        ws = wb.create_sheet("File Shares")
        
        headers = ['Name', 'Share Path', 'Total Size (GB)', 'Share Type', 'Environment', 'Status']
        for i, header in enumerate(headers, 1):
            ws.cell(row=1, column=i, value=header).font = Font(bold=True)
        
        FileShare = self.models.get('FileShare')
        if FileShare:
            file_shares = FileShare.query.all()
            for i, fs in enumerate(file_shares, 2):
                ws.cell(row=i, column=1, value=fs.name)
                ws.cell(row=i, column=2, value=fs.share_path)
                ws.cell(row=i, column=3, value=fs.total_size_gb)
                ws.cell(row=i, column=4, value=fs.share_type)
                ws.cell(row=i, column=5, value=fs.environment)
                ws.cell(row=i, column=6, value=fs.status)
    
    def _create_cost_analysis_sheet(self, wb):
        """Create cost analysis sheet"""
        ws = wb.create_sheet("Cost Analysis")
        
        ws['A1'] = "Cost Analysis Summary"
        ws['A1'].font = Font(bold=True, size=14)
        
        cost_data = self._get_cost_summary()
        
        ws['A3'] = "Cost Category"
        ws['B3'] = "Amount (USD)"
        
        cost_items = [
            ('Annual Cloud Infrastructure', cost_data['annual_cloud_cost']),
            ('Migration Services (One-time)', cost_data['migration_services_cost']),
            ('Training & Support', cost_data.get('training_cost', 5000)),
            ('Total First Year Cost', cost_data['total_first_year'])
        ]
        
        for i, (category, amount) in enumerate(cost_items, 4):
            ws.cell(row=i, column=1, value=category)
            ws.cell(row=i, column=2, value=amount)
    
    def _create_timeline_sheet(self, wb):
        """Create timeline sheet"""
        ws = wb.create_sheet("Timeline")
        
        ws['A1'] = "Migration Timeline"
        ws['A1'].font = Font(bold=True, size=14)
        
        timeline_data = self._get_timeline_summary()
        
        ws['A3'] = "Phase"
        ws['B3'] = "Description"
        ws['C3'] = "Duration (weeks)"
        ws['D3'] = "Start Week"
        ws['E3'] = "End Week"
        
        phases = [
            ('Assessment & Planning', 'Inventory assessment and migration planning', 4, 1, 4),
            ('Environment Setup', 'Cloud infrastructure setup', 3, 5, 7),
            ('Data Migration', 'Database and file share migration', 6, 8, 13),
            ('Server Migration', 'Application and server migration', 3, 14, 16)
        ]
        
        for i, (phase, desc, duration, start, end) in enumerate(phases, 4):
            ws.cell(row=i, column=1, value=phase)
            ws.cell(row=i, column=2, value=desc)
            ws.cell(row=i, column=3, value=duration)
            ws.cell(row=i, column=4, value=start)
            ws.cell(row=i, column=5, value=end)
    
    def _get_summary_data(self):
        """Get comprehensive summary data for reports"""
        Server = self.models.get('Server')
        Database = self.models.get('Database')
        FileShare = self.models.get('FileShare')
        
        servers_count = Server.query.count() if Server else 0
        databases_count = Database.query.count() if Database else 0
        file_shares_count = FileShare.query.count() if FileShare else 0
        
        # Calculate totals from servers
        total_vcpu = 0
        total_ram_gb = 0
        total_disk_gb = 0
        
        if Server:
            try:
                total_vcpu = self.db.session.query(self.db.func.sum(Server.vcpu)).scalar() or 0
                total_ram_gb = self.db.session.query(self.db.func.sum(Server.ram)).scalar() or 0
                total_disk_gb = self.db.session.query(self.db.func.sum(Server.disk_size)).scalar() or 0
            except:
                servers = Server.query.all()
                for server in servers:
                    total_vcpu += getattr(server, 'vcpu', 0) or 0
                    total_ram_gb += getattr(server, 'ram', 0) or 0
                    total_disk_gb += getattr(server, 'disk_size', 0) or 0
        
        # Calculate database sizes
        total_db_size = 0
        if Database:
            try:
                total_db_size = self.db.session.query(self.db.func.sum(Database.size_gb)).scalar() or 0
            except:
                databases = Database.query.all()
                for db in databases:
                    total_db_size += getattr(db, 'size_gb', 0) or 0
                
        # Calculate file share sizes
        total_fs_size = 0
        if FileShare:
            try:
                total_fs_size = self.db.session.query(self.db.func.sum(FileShare.total_size_gb)).scalar() or 0
            except:
                file_shares = FileShare.query.all()
                for fs in file_shares:
                    total_fs_size += getattr(fs, 'total_size_gb', 0) or 0
        
        # Calculate complexity metrics
        total_components = servers_count + databases_count + file_shares_count
        complexity_score = min(10, max(1, (servers_count + databases_count * 1.5 + file_shares_count * 0.5) / 5))
        complexity_level = "Low" if complexity_score < 3 else "Medium" if complexity_score < 7 else "High"
        
        # Estimate costs (basic calculations)
        monthly_cost = (servers_count * 200) + (databases_count * 300) + ((total_db_size + total_fs_size) * 0.1)
        annual_cost = monthly_cost * 12
        
        # Estimate timeline
        estimated_weeks = max(8, 4 + (servers_count * 1.5) + (databases_count * 2) + (file_shares_count * 0.5))
        
        return {
            'servers_count': servers_count,
            'databases_count': databases_count,
            'file_shares_count': file_shares_count,
            'total_vcpu': total_vcpu,
            'total_ram_gb': total_ram_gb,
            'total_disk_gb': total_disk_gb,
            'total_db_size': total_db_size,
            'total_fs_size': total_fs_size,
            'total_data_gb': total_disk_gb + total_db_size + total_fs_size,
            'total_components': total_components,
            'complexity_score': complexity_score,
            'complexity_level': complexity_level,
            'estimated_weeks': int(estimated_weeks),
            'monthly_cost': monthly_cost,
            'annual_cost': annual_cost,
            'primary_strategy': 'Rehost (Lift & Shift)'
        }
    
    def _get_cost_summary(self):
        """Get cost summary data"""
        summary_data = self._get_summary_data()
        
        # Basic cost calculations based on inventory
        server_cost = summary_data['servers_count'] * 100 * 12  # $100/month per server
        database_cost = summary_data['databases_count'] * 200 * 12  # $200/month per database
        storage_cost = summary_data['total_data_gb'] * 0.10 * 12  # $0.10/GB per month
        
        annual_cloud_cost = server_cost + database_cost + storage_cost
        migration_services_cost = summary_data['servers_count'] * 500 + summary_data['databases_count'] * 1000
        
        return {
            'annual_cloud_cost': annual_cloud_cost,
            'migration_services_cost': migration_services_cost,
            'total_first_year': annual_cloud_cost + migration_services_cost,
            'training_cost': 5000
        }
    
    def _get_timeline_summary(self):
        """Get timeline summary data"""
        summary_data = self._get_summary_data()
        
        # Calculate timeline based on inventory
        base_weeks = 8
        additional_weeks = (summary_data['servers_count'] // 10) + (summary_data['databases_count'] // 5)
        total_weeks = base_weeks + additional_weeks
        
        return {
            'total_weeks': total_weeks,
            'end_date': 'Based on selected start date',
            'phases': 4,
            'critical_path': ['Assessment', 'Data Migration', 'Server Migration']
        }
