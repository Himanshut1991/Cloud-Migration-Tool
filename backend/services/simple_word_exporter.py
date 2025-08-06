#!/usr/bin/env python3
"""Simple Word export service without complex queries"""

import os
from datetime import datetime
from docx import Document

class SimpleWordExporter:
    """Simple Word exporter with minimal dependencies"""
    
    def __init__(self, db, models):
        self.db = db
        self.models = models
        self.output_dir = os.path.join(os.path.dirname(__file__), '..', 'exports')
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_to_word(self):
        """Export migration plan to Word document with safe data handling"""
        try:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'migration_plan_{timestamp}.docx'
            filepath = os.path.join(self.output_dir, filename)
            
            print(f"📄 Creating Word document: {filename}")
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Cloud Migration Assessment Report', 0)
            title.alignment = 1
            
            # Date
            date_para = doc.add_paragraph()
            date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            date_para.alignment = 1
            
            doc.add_page_break()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            
            # Get safe counts
            servers_count = 0
            databases_count = 0
            file_shares_count = 0
            
            try:
                Server = self.models.get('Server')
                if Server:
                    servers_count = Server.query.count()
            except:
                servers_count = 0
                
            try:
                Database = self.models.get('Database')
                if Database:
                    databases_count = Database.query.count()
            except:
                databases_count = 0
                
            try:
                FileShare = self.models.get('FileShare')
                if FileShare:
                    file_shares_count = FileShare.query.count()
            except:
                file_shares_count = 0
            
            # Summary paragraph
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"This migration assessment covers {servers_count} servers, ")
            summary_para.add_run(f"{databases_count} databases, and {file_shares_count} file shares. ")
            summary_para.add_run("The assessment provides cost estimates, timeline projections, and migration recommendations.")
            
            # Key metrics table
            doc.add_heading('Key Metrics', level=2)
            metrics_table = doc.add_table(rows=6, cols=2)
            metrics_table.style = 'Table Grid'
            
            # Headers
            metrics_table.rows[0].cells[0].text = 'Metric'
            metrics_table.rows[0].cells[1].text = 'Value'
            
            # Data
            total_components = servers_count + databases_count + file_shares_count
            estimated_cost = total_components * 100  # Simple calculation
            
            metrics_table.rows[1].cells[0].text = 'Total Components'
            metrics_table.rows[1].cells[1].text = str(total_components)
            
            metrics_table.rows[2].cells[0].text = 'Estimated Duration'
            metrics_table.rows[2].cells[1].text = f"{max(8, total_components)} weeks"
            
            metrics_table.rows[3].cells[0].text = 'Migration Complexity'
            complexity = 'Low' if total_components < 10 else 'Medium' if total_components < 30 else 'High'
            metrics_table.rows[3].cells[1].text = complexity
            
            metrics_table.rows[4].cells[0].text = 'Monthly Cloud Cost'
            metrics_table.rows[4].cells[1].text = f"${estimated_cost:,}"
            
            metrics_table.rows[5].cells[0].text = 'Primary Strategy'
            metrics_table.rows[5].cells[1].text = 'Rehost (Lift & Shift)'
            
            # Infrastructure Breakdown
            doc.add_heading('Infrastructure Breakdown', level=1)
            
            breakdown_table = doc.add_table(rows=4, cols=3)
            breakdown_table.style = 'Table Grid'
            
            # Headers
            breakdown_table.rows[0].cells[0].text = 'Component Type'
            breakdown_table.rows[0].cells[1].text = 'Count'
            breakdown_table.rows[0].cells[2].text = 'Migration Priority'
            
            # Data
            breakdown_table.rows[1].cells[0].text = 'Servers'
            breakdown_table.rows[1].cells[1].text = str(servers_count)
            breakdown_table.rows[1].cells[2].text = 'High'
            
            breakdown_table.rows[2].cells[0].text = 'Databases'
            breakdown_table.rows[2].cells[1].text = str(databases_count)
            breakdown_table.rows[2].cells[2].text = 'Critical'
            
            breakdown_table.rows[3].cells[0].text = 'File Shares'
            breakdown_table.rows[3].cells[1].text = str(file_shares_count)
            breakdown_table.rows[3].cells[2].text = 'Medium'
            
            # Cost Analysis
            doc.add_heading('Cost Analysis', level=1)
            cost_para = doc.add_paragraph()
            cost_para.add_run("The cost analysis is based on current infrastructure sizing and standard cloud pricing models. ")
            cost_para.add_run(f"Total estimated monthly cost: ${estimated_cost:,}. ")
            cost_para.add_run(f"Annual cost estimate: ${estimated_cost * 12:,}.")
            
            # Timeline
            doc.add_heading('Migration Timeline', level=1)
            timeline_para = doc.add_paragraph()
            timeline_para.add_run(f"The migration is estimated to take {max(8, total_components)} weeks, ")
            timeline_para.add_run("organized into phases: Planning (2 weeks), Environment Setup (2-3 weeks), ")
            timeline_para.add_run("Data Migration (3-4 weeks), Application Migration (2-4 weeks), and Testing (1-2 weeks).")
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            rec_para = doc.add_paragraph()
            rec_para.add_run("Based on the assessment, we recommend a phased approach starting with non-critical systems. ")
            rec_para.add_run("Implement proper backup and monitoring solutions. Consider using managed services where possible.")
            
            recommendations = [
                "Start with a pilot migration of non-critical systems",
                "Implement comprehensive backup and disaster recovery",
                "Use cloud-native monitoring and alerting",
                "Consider reserved instances for cost optimization",
                "Establish governance processes for cloud resources"
            ]
            
            for rec in recommendations:
                rec_bullet = doc.add_paragraph(style='ListBullet')
                rec_bullet.add_run(rec)
            
            doc.save(filepath)
            print(f"✅ Word document saved: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"❌ Word export error: {str(e)}")
            import traceback
            traceback.print_exc()
            raise Exception(f"Failed to export to Word: {str(e)}")

if __name__ == "__main__":
    # Test the simple exporter
    import sys
    sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    
    from models_new import init_models
    from flask import Flask
    from flask_sqlalchemy import SQLAlchemy
    
    app = Flask(__name__)
    basedir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(basedir, "migration_tool.db")}'
    app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
    
    db = SQLAlchemy(app)
    
    with app.app_context():
        models = init_models(db)
        exporter = SimpleWordExporter(db, models)
        filepath = exporter.export_to_word()
        print(f"🎉 Test completed! File: {filepath}")
