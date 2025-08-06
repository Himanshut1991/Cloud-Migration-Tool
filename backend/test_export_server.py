#!/usr/bin/env python3
"""Minimal export test server"""

from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import os
import json
from datetime import datetime

app = Flask(__name__)
CORS(app)

exports_dir = os.path.join(os.path.dirname(__file__), 'exports')
os.makedirs(exports_dir, exist_ok=True)

@app.route('/api/export', methods=['POST'])
def test_export():
    """Test export endpoint"""
    try:
        data = request.get_json() or {}
        format_type = data.get('format', 'pdf').lower()
        
        print(f"🧪 Testing {format_type} export...")
        
        if format_type == 'pdf':
            # Create a simple test PDF
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'test_migration_plan_{timestamp}.pdf'
            filepath = os.path.join(exports_dir, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            story.append(Paragraph("Cloud Migration Test Report", styles['Title']))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Paragraph("This is a test PDF export to verify the system is working.", styles['Normal']))
            
            doc.build(story)
            
        elif format_type == 'excel':
            # Create a simple test Excel file
            from openpyxl import Workbook
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'test_migration_plan_{timestamp}.xlsx'
            filepath = os.path.join(exports_dir, filename)
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Test Report"
            
            ws['A1'] = "Cloud Migration Test Report"
            ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            ws['A4'] = "Component"
            ws['B4'] = "Count"
            ws['A5'] = "Servers"
            ws['B5'] = 10
            ws['A6'] = "Databases"
            ws['B6'] = 5
            
            wb.save(filepath)
            
        elif format_type == 'word':
            # Create a simple test Word document
            from docx import Document
            
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'test_migration_plan_{timestamp}.docx'
            filepath = os.path.join(exports_dir, filename)
            
            doc = Document()
            doc.add_heading('Cloud Migration Test Report', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("This is a test Word export to verify the system is working.")
            
            table = doc.add_table(rows=3, cols=2)
            table.rows[0].cells[0].text = 'Component'
            table.rows[0].cells[1].text = 'Count'
            table.rows[1].cells[0].text = 'Servers'
            table.rows[1].cells[1].text = '10'
            table.rows[2].cells[0].text = 'Databases'
            table.rows[2].cells[1].text = '5'
            
            doc.save(filepath)
            
        else:
            return jsonify({'error': f'Unsupported format: {format_type}'}), 400
        
        # Get file info
        file_size = os.path.getsize(filepath)
        
        print(f"✅ {format_type} export completed: {filename}")
        
        return jsonify({
            'message': f'{format_type.upper()} export completed successfully',
            'format': format_type,
            'filename': filename,
            'filepath': filepath,
            'file_size': file_size,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        print(f"❌ Export failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Export failed: {str(e)}'}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_test_file(filename):
    """Download test file"""
    try:
        filepath = os.path.join(exports_dir, filename)
        if os.path.exists(filepath):
            return send_file(filepath, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/dashboard', methods=['GET'])
def test_dashboard():
    """Test dashboard endpoint"""
    return jsonify({
        'message': 'Test backend is running',
        'timestamp': datetime.now().isoformat(),
        'exports_available': ['pdf', 'excel', 'word']
    })

if __name__ == '__main__':
    print("🧪 Starting test export server on http://localhost:5001")
    print("Available endpoints:")
    print("  - GET  /api/dashboard")
    print("  - POST /api/export")
    print("  - GET  /api/download/<filename>")
    app.run(host='0.0.0.0', port=5001, debug=True)
