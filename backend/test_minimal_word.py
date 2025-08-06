#!/usr/bin/env python3
"""Minimal Word export test to isolate the error"""

import os
import sys
sys.path.append('.')

from docx import Document
from services.export_service_new import ExportService
from models_new import init_models
from flask import Flask
from flask_sqlalchemy import SQLAlchemy

def test_minimal_word():
    """Test minimal Word creation"""
    try:
        print('🧪 Testing minimal Word document creation...')
        
        # Create a simple document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        
        # Add a simple table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = 'Header 1'
        table.rows[0].cells[1].text = 'Header 2'
        table.rows[1].cells[0].text = 'Data 1'
        table.rows[1].cells[1].text = 'Data 2'
        
        test_file = 'test_minimal.docx'
        doc.save(test_file)
        
        print(f'✅ Minimal Word document created: {test_file}')
        
        # Now test with our export service
        print('🧪 Testing with export service...')
        
        app = Flask(__name__)
        basedir = os.path.dirname(os.path.abspath(__file__))
        app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(basedir, "migration_tool.db")}'
        app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
        
        db = SQLAlchemy(app)
        
        with app.app_context():
            models = init_models(db)
            export_service = ExportService(db, models)
            
            # Get summary data to check for issues
            print('🔍 Getting summary data...')
            summary_data = export_service._get_summary_data()
            print(f'Summary data keys: {list(summary_data.keys())}')
            print(f'Summary data values: {summary_data}')
            
            # Test creating a Word document with our service
            print('🧪 Testing export service Word creation...')
            word_path = export_service.export_to_word()
            print(f'✅ Export service Word document created: {word_path}')
            
    except Exception as e:
        print(f'❌ Error in test: {e}')
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_minimal_word()
